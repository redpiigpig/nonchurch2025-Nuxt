# -*- coding: utf-8 -*-
"""
Apply all DB corrections to issue 7 Word documents.
Uses run.text replacement to preserve fonts and formatting.
"""

import sys
import os
from docx import Document
from lxml import etree
import copy

BASE = r"c:\Users\user\Desktop\nonchurch-nuxt\stores\無境界者雜誌\07-第七期"

# ────────────────────────────────────────────────────────────────────────────
# Corrections: (file_key, old_text, new_text)
# file_key maps to filename
# ────────────────────────────────────────────────────────────────────────────
FILES = {
    "7-1":  "7-1編輯室報告.docx",
    "7-2":  "7-2本期作者簡介.docx",
    "7-3":  "7-3火燒島的美麗風景.docx",
    "7-4":  "7-4教會外的福音見證者.docx",
    "7-5":  "7-5聆聽被遺忘的苦難.docx",
    "7-6":  "7-6In是Siáng？.docx",
    "7-7":  "7-7被國家剝奪權利的人.docx",
    "7-8":  "7-8黑暗中的那點光.docx",
    "7-9":  "7-9跨地域的緣與情.docx",
    "7-10": "7-10帝大教授的殖民地診斷書.docx",
    "7-11": "7-11戴著鎖鏈歡呼的先知.docx",
    "7-12": "7-12願月光指引我們的故事.docx",
    "7-13": "7-13在繪本裡重讀神聖.docx",
    "7-14": "7-14蔡依林有沒有在拜撒旦到底干那些基督徒什麼事？.docx",
    "7-15": "7-15從智慧到十字架.docx",
    "7-16": "7-16基督宗教的七個大公傳統.docx",
}

CORRECTIONS = [
    # ── 7-1 ──────────────────────────────────────────────────────────────
    ("7-1", "藉著了介紹",       "藉著介紹"),
    ("7-1", "此以為核心",       "以此為核心"),
    ("7-1", "要論說我們的生命", "訴說我們的生命"),

    # ── 7-2 ──────────────────────────────────────────────────────────────
    ("7-2", "尋找着",           "尋找著"),
    ("7-2", "1967以僑生身分",   "1967年以僑生身分"),

    # ── 7-5 ──────────────────────────────────────────────────────────────
    ("7-5", "生出面對苦難",     "孕育出面對苦難"),
    ("7-5", "去成都還不知道",   "去成，都還不知道"),
    ("7-5", "與和丈夫同一國",   "和丈夫同一國"),
    ("7-5", "西伯利亞戰友們",   "西伯利亞的戰友們"),

    # ── 7-6 ──────────────────────────────────────────────────────────────
    ("7-6", "讀者讀者",         "讀者"),

    # ── 7-7 ──────────────────────────────────────────────────────────────
    ("7-7", "Just-world fallacy",   "Just-world theory"),
    ("7-7", "被一同被押上車",       "被一同押上車"),
    ("7-7", "積極的活出",           "積極地活出"),
    ("7-7", "直接了當地",           "直截了當地"),

    # ── 7-8 ──────────────────────────────────────────────────────────────
    ("7-8", "他們願意陪著",     "他/她願意陪著"),
    ("7-8", "沈默",             "沉默"),
    ("7-8", "沈重",             "沉重"),

    # ── 7-9 ──────────────────────────────────────────────────────────────
    ("7-9", "是指著矢内原",     "是指矢内原"),
    ("7-9", "留下印象的是，則是", "留下印象的，則是"),
    ("7-9", "不僅陳茂棠",       "不僅是陳茂棠"),
    ("7-9", "忽然之間會固定早退", "忽然開始固定早退"),
    ("7-9", "規劃名為",         "歸化名為"),
    ("7-9", "被譽爲",           "被譽為"),
    ("7-9", "進而，以下將談及", "接下來，以下將談及"),
    ("7-9", "足立並新渡戸",     "足立與新渡戶"),
    ("7-9", "Yanaibara",        "Yanaihara"),
    ("7-9", "Fujjida",          "Fujita"),
    ("7-9", "統治時期間",       "統治時期"),

    # ── 7-10 ─────────────────────────────────────────────────────────────
    ("7-10", "滿州",    "滿洲"),
    ("7-10", "具体的",  "具體的"),
    ("7-10", "在于",    "在於"),

    # ── 7-11 ─────────────────────────────────────────────────────────────
    ("7-11", "但他不行",          "但他無法如願"),
    ("7-11", "斯哥德爾摩",        "斯德哥爾摩"),
    ("7-11", "幾盡全力",          "竭盡全力"),
    ("7-11", "被缚",              "被縛"),
    ("7-11", "一般情侶中的角度",  "一般情侶的角度"),
    ("7-11", "托管",              "託管"),
    ("7-11", "猶太國",            "猶大國"),

    # ── 7-12 ─────────────────────────────────────────────────────────────
    ("7-12", "荒繆", "荒謬"),

    # ── 7-13 ─────────────────────────────────────────────────────────────
    ("7-13", "生態心理學",  "精神分析學派"),
    ("7-13", "護定環境",    "抱持環境"),

    # ── 7-14 ─────────────────────────────────────────────────────────────
    ("7-14", "主辦方的口供不一致", "主辦方的說法不一致"),
    ("7-14", "人之常性",          "人之常情"),
    # The "整理其教義*成" → "並整理其教義，將其化為" was already applied in DB;
    # in the Word doc the raw text is likely: 整理其教義成...
    # We'll try the raw-text version:
    ("7-14", "並整理其教義成六大提倡", "並整理其教義，將其化為六大提倡"),

    # ── 7-16 ─────────────────────────────────────────────────────────────
    ("7-16", "並其視為普世教義",  "並將其視為普世教義"),
    ("7-16", "在元325年",         "公元325年"),
    ("7-16", "主張只能其為",      "主張只能稱其為"),
    ("7-16", "狄奧多庫若",        "狄奧斯庫若"),
    ("7-16", "首都埃奇米亞津",    "聖城埃奇米亞津"),
]

# ────────────────────────────────────────────────────────────────────────────
# Footnote corrections: (file_key, fn_index_1based, old_text, new_text)
# ────────────────────────────────────────────────────────────────────────────
FOOTNOTE_CORRECTIONS = [
    ("7-7",  8, "以預防他如果在大馬被逮捕的話", "以防他若在大馬被逮捕"),
    ("7-7",  9, "J. Glenn Beall. Jr",           "J. Glenn Beall, Jr."),
    ("7-11", 4, "一切事務",                      "一切事物"),
]


def get_para_full_text(para):
    return "".join(r.text for r in para.runs)


def replace_in_para(para, old, new):
    """
    Try to replace 'old' with 'new' within a paragraph.
    Strategy:
      1. If any single run contains 'old', replace in that run (format safe).
      2. Otherwise concatenate adjacent runs, find the span, redistribute text
         while preserving the first run's formatting for the changed portion.
    Returns True if a replacement was made.
    """
    # Strategy 1: single-run replacement
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True

    # Strategy 2: multi-run merge
    full = get_para_full_text(para)
    if old not in full:
        return False

    # Find position in concatenated text
    idx = full.index(old)
    end = idx + len(old)

    # Identify which runs are affected
    pos = 0
    runs = para.runs
    affected_start = None
    affected_end = None
    run_spans = []
    for i, run in enumerate(runs):
        run_start = pos
        run_end = pos + len(run.text)
        run_spans.append((run_start, run_end))
        if run_start <= idx < run_end and affected_start is None:
            affected_start = i
        if run_start < end <= run_end and affected_end is None:
            affected_end = i
        pos = run_end

    if affected_start is None or affected_end is None:
        return False

    # Rebuild: put new text in first affected run, clear the rest
    pre  = full[run_spans[affected_start][0]:idx]
    post = full[end:run_spans[affected_end][1]]
    runs[affected_start].text = pre + new + post
    for i in range(affected_start + 1, affected_end + 1):
        runs[i].text = ""
    return True


def replace_in_doc_body(doc, old, new):
    count = 0
    for para in doc.paragraphs:
        if replace_in_para(para, old, new):
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replace_in_para(para, old, new):
                        count += 1
    return count


def replace_in_footnotes(doc, fn_number, old, new):
    """
    Access native Word footnotes via XML.
    Footnotes are stored in the footnotes part.
    fn_number is 1-based (matches [1], [2], etc. in text).
    Returns count of replacements made.
    """
    try:
        fn_part = doc.part.footnotes
    except Exception:
        return 0

    if fn_part is None:
        return 0

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    root = fn_part._element

    count = 0
    # Footnotes in OOXML have w:footnote elements with w:id attribute
    # id=0 is separator, id=1 is first user footnote, etc.
    for fn_el in root.findall('.//w:footnote', ns):
        fn_id = fn_el.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
        if fn_id is None:
            continue
        try:
            fn_id_int = int(fn_id)
        except ValueError:
            continue
        # id 0 and -1 are reserved; user footnotes start at 1
        if fn_id_int != fn_number:
            continue

        # Collect all w:r (runs) within this footnote
        for r_el in fn_el.findall('.//w:r', ns):
            t_els = r_el.findall('w:t', ns)
            for t_el in t_els:
                if t_el.text and old in t_el.text:
                    t_el.text = t_el.text.replace(old, new)
                    count += 1

        # Multi-run merge within this footnote paragraph
        if count == 0:
            for p_el in fn_el.findall('.//w:p', ns):
                runs_t = [r_el.findall('w:t', ns) for r_el in p_el.findall('.//w:r', ns)]
                runs_flat = p_el.findall('.//w:r', ns)
                full = ""
                spans = []
                for r_el in runs_flat:
                    start = len(full)
                    text = "".join((t.text or "") for t in r_el.findall('w:t', ns))
                    full += text
                    spans.append((start, len(full), r_el))

                if old not in full:
                    continue

                idx = full.index(old)
                end_idx = idx + len(old)
                a_start = next((i for i, (s, e, _) in enumerate(spans) if s <= idx < e), None)
                a_end   = next((i for i, (s, e, _) in enumerate(spans) if s < end_idx <= e), None)

                if a_start is None or a_end is None:
                    continue

                # Merge text into first affected run
                pre  = full[spans[a_start][0]:idx]
                post = full[end_idx:spans[a_end][1]]
                new_text = pre + new + post

                # Update first run's w:t
                first_r = spans[a_start][2]
                t_list = first_r.findall('w:t', ns)
                if t_list:
                    t_list[0].text = new_text
                    for t in t_list[1:]:
                        t.text = ""
                else:
                    t_new = etree.SubElement(first_r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                    t_new.text = new_text

                # Clear subsequent runs
                for i in range(a_start + 1, a_end + 1):
                    for t in spans[i][2].findall('w:t', ns):
                        t.text = ""
                count += 1

    return count


def main():
    # Group corrections by file
    by_file = {}
    for (key, old, new) in CORRECTIONS:
        by_file.setdefault(key, []).append((old, new))

    fn_by_file = {}
    for (key, fn_num, old, new) in FOOTNOTE_CORRECTIONS:
        fn_by_file.setdefault(key, []).append((fn_num, old, new))

    all_keys = sorted(set(list(by_file.keys()) + list(fn_by_file.keys())))

    for key in all_keys:
        fname = FILES.get(key)
        if not fname:
            print(f"[SKIP] {key}: no filename mapping")
            continue

        fpath = os.path.join(BASE, fname)
        if not os.path.exists(fpath):
            print(f"[SKIP] {key}: file not found — {fpath}")
            continue

        doc = Document(fpath)
        changed = False
        print(f"\n{'='*60}")
        print(f"  {key}: {fname}")
        print(f"{'='*60}")

        for (old, new) in by_file.get(key, []):
            n = replace_in_doc_body(doc, old, new)
            if n > 0:
                print(f"  OK [{n}x] {old!r} → {new!r}")
                changed = True
            else:
                print(f"  !!  NOT FOUND: {old!r}")

        for (fn_num, old, new) in fn_by_file.get(key, []):
            n = replace_in_footnotes(doc, fn_num, old, new)
            if n > 0:
                print(f"  OK fn{fn_num} [{n}x] {old!r} → {new!r}")
                changed = True
            else:
                print(f"  !!  fn{fn_num} NOT FOUND: {old!r}")

        if changed:
            doc.save(fpath)
            print(f"  >> Saved.")
        else:
            print(f"  (no changes)")

    print("\nDone.")


if __name__ == "__main__":
    main()
