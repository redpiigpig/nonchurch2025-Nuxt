#!/usr/bin/env python3
"""
patch_docx.py  ──  兩步驟 Word 排版修補工具

━━ 步驟一：產生螢光筆審稿檔 ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  python scripts/patch_docx.py review <期數> [Word資料夾]

  Word資料夾可省略：預設自動找 G:/我的雲端硬碟/資料/無境界者/雜誌/0{N}-第{N}期/

  例：
    python scripts/patch_docx.py review 7
    python scripts/patch_docx.py review 7 D:/其他路徑/

  輸出（輸出到 Word 資料夾旁）：
    review_issue7.docx   ← 合併稿，黃螢光筆 = 需改動，灰字提示新版本
    review_issue7.json   ← 改動清單（供第二步驟使用）

━━ 步驟二：套用已審核的改動 ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  python scripts/patch_docx.py apply <review_doc.docx> [Word資料夾] [--inplace]

  例：
    python scripts/patch_docx.py apply review_issue7.docx
    python scripts/patch_docx.py apply review_issue7.docx D:/其他路徑/ --inplace

  --inplace：直接覆蓋原檔（預設輸出 _patched.docx）

  邏輯：
    · 讀 review_issue7.json 知道哪些段落有改動
    · 在 review_issue7.docx 中仍有黃螢光筆 → 使用者同意 → 套用
    · 螢光筆已移除 → 使用者拒絕 → 跳過
    · 套用時只改對應 run 的文字，字型/大小/顏色完整保留

━━ 依賴套件（pip install） ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  requests  python-docx  beautifulsoup4  python-dotenv（可選）
"""

import sys
import os
import re
import json
import difflib
import io
from copy import deepcopy
from pathlib import Path

# Windows 終端機預設 cp950，強制改為 utf-8 輸出
if hasattr(sys.stdout, "buffer"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "buffer"):
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

# ── 依賴檢查 ──────────────────────────────────────────────────────────────────
def _require(pip_name, import_name=None):
    import importlib
    try:
        return importlib.import_module(import_name or pip_name)
    except ImportError:
        print(f"[X]  缺少套件，請先執行：pip install {pip_name}", file=sys.stderr)
        sys.exit(1)

requests_mod = _require("requests")
docx_mod     = _require("python-docx", "docx")
bs4_mod      = _require("beautifulsoup4", "bs4")

Document      = docx_mod.Document
BeautifulSoup = bs4_mod.BeautifulSoup
from docx.oxml.ns   import qn
from docx.shared    import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX

# ── 載入 .env ─────────────────────────────────────────────────────────────────
def _load_env():
    env_path = Path(__file__).parent.parent / ".env"
    if not env_path.exists():
        return
    try:
        from dotenv import load_dotenv
        load_dotenv(env_path)
        return
    except ImportError:
        pass
    with open(env_path, encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, _, v = line.partition("=")
                os.environ.setdefault(k.strip(), v.strip().strip('"').strip("'"))

_load_env()

# ── 自動找期數資料夾 ──────────────────────────────────────────────────────────
_STORES_BASE = Path(__file__).parent.parent / "stores" / "無境界者雜誌"

def auto_issue_folder(issue_number):
    """自動找 G:/我的雲端硬碟/資料/無境界者/雜誌/{NN}-* 資料夾"""
    if not _STORES_BASE.is_dir():
        return None
    prefix = f"{int(issue_number):02d}-"
    matches = [d for d in _STORES_BASE.iterdir()
               if d.is_dir() and d.name.startswith(prefix)]
    return matches[0] if matches else None

# ── Supabase ──────────────────────────────────────────────────────────────────
def _supabase_cfg():
    base = os.environ.get("VITE_SUPABASE_URL", "").rstrip("/")
    key  = os.environ.get("SUPABASE_SECRET_KEY") or os.environ.get("VITE_SUPABASE_KEY")
    missing = []
    if not base: missing.append("VITE_SUPABASE_URL")
    if not key:  missing.append("SUPABASE_SECRET_KEY")
    if missing:
        print(f"❌  .env 缺少：{', '.join(missing)}", file=sys.stderr)
        sys.exit(1)
    return base, {"apikey": key, "Authorization": f"Bearer {key}"}

def fetch_issue_articles(issue_number):
    base, headers = _supabase_cfg()
    url = (f"{base}/rest/v1/articles"
           f"?issue=eq.{issue_number}"
           f"&select=id,title,content"
           f"&order=sort_order.asc")
    resp = requests_mod.get(url, headers=headers, timeout=20)
    if not resp.ok:
        print(f"❌  Supabase 查詢失敗 ({resp.status_code}): {resp.text}", file=sys.stderr)
        sys.exit(1)
    data = resp.json()
    if not data:
        print(f"❌  找不到第 {issue_number} 期文章", file=sys.stderr)
        sys.exit(1)
    return data

# ── HTML → 純文字段落清單 ──────────────────────────────────────────────────────
_FN_RE = re.compile(
    r'\[\^\d+\]|<sup[^>]*class=["\']footnote-ref["\'][^>]*>.*?</sup>',
    re.IGNORECASE | re.DOTALL
)

def _strip_html(html):
    text = _FN_RE.sub("", html or "")
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    import html as _h
    return _h.unescape(text).strip()

def extract_html_paras(html_content):
    """HTML 內容 → 純文字段落清單"""
    if not html_content:
        return []
    soup = BeautifulSoup(html_content, "html.parser")
    results = []

    def _visit(el):
        if not hasattr(el, "name") or not el.name:
            return
        if el.name in ("p", "h2", "h3", "blockquote", "li"):
            t = _strip_html(str(el))
            if t:
                results.append(t)
        else:
            for child in el.children:
                _visit(child)

    for child in soup.children:
        _visit(child)
    return results

# ── Word 段落工具 ──────────────────────────────────────────────────────────────
def para_text(para):
    return "".join(r.text for r in para.runs)

def collect_nonempty_paras(doc):
    """回傳 [(原始索引, para, text), ...]，僅含有文字段落"""
    return [(i, p, para_text(p))
            for i, p in enumerate(doc.paragraphs)
            if para_text(p).strip()]

# ── 段落對齊 ──────────────────────────────────────────────────────────────────
def align_paras(word_entries, html_texts):
    """
    回傳 {word_orig_idx: {'h_text': ..., 'w_text': ..., 'sim': ...}}
    只含有差異且相似度 >= 0.5 的配對。
    """
    w_texts = [t for _, _, t in word_entries]
    sm = difflib.SequenceMatcher(None, w_texts, list(html_texts), autojunk=False)
    result = {}
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "replace":
            for off in range(min(i2-i1, j2-j1)):
                wt = w_texts[i1+off]
                ht = html_texts[j1+off]
                sim = difflib.SequenceMatcher(None, wt, ht).ratio()
                if sim >= 0.5 and wt != ht:
                    w_orig_idx = word_entries[i1+off][0]
                    result[w_orig_idx] = {"h_text": ht, "w_text": wt, "sim": sim}
    return result

# ── 字符級 diff & run-level 修補 ─────────────────────────────────────────────
def _run_map(runs):
    rm, pos = [], 0
    for i, r in enumerate(runs):
        rm.append((pos, pos + len(r.text), i))
        pos += len(r.text)
    return rm

def try_patch_para(para, new_text):
    """
    嘗試把 para 的文字更新為 new_text，保留 run 格式。
    回傳 (True, msg) 或 (False, reason)。
    """
    runs = para.runs
    if not runs:
        return False, "段落無 runs"
    old_text = "".join(r.text for r in runs)
    if old_text == new_text:
        return True, "無差異"

    rm = _run_map(runs)
    total = len(old_text)
    sm = difflib.SequenceMatcher(None, old_text, new_text, autojunk=False)
    ops = [(t, i1, i2, j1, j2)
           for t, i1, i2, j1, j2 in sm.get_opcodes() if t != "equal"]
    if not ops:
        return True, "無差異"

    patches = []
    for tag, i1, i2, j1, j2 in ops:
        seg = new_text[j1:j2]
        found = False
        for rs, re_, ri in rm:
            if tag == "insert":
                if rs <= i1 <= re_:
                    patches.append((ri, i1-rs, i1-rs, seg))
                    found = True
                    break
            else:
                if rs <= i1 and i2 <= re_:
                    patches.append((ri, i1-rs, i2-rs, seg))
                    found = True
                    break
        if not found:
            if tag == "insert" and i1 == total:
                last_rs, last_re, last_ri = rm[-1]
                patches.append((last_ri, last_re-last_rs, last_re-last_rs, seg))
            else:
                snip_old = repr(old_text[max(0, i1-8):i2+8])
                snip_new = repr(new_text[max(0, j1-8):j2+8])
                return False, (f"改動跨越多個 run\n"
                               f"      Word：{snip_old}\n"
                               f"      新版：{snip_new}")

    patches.sort(key=lambda x: (x[0], x[1]), reverse=True)
    for ri, ps, pe, seg in patches:
        r = runs[ri]
        r.text = r.text[:ps] + seg + r.text[pe:]
    return True, f"套用 {len(ops)} 處改動"

# ── 計算需標示的字符範圍（在 old_text 中） ────────────────────────────────────
def changed_ranges_in_old(old_text, new_text):
    sm = difflib.SequenceMatcher(None, old_text, new_text, autojunk=False)
    ranges = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag in ("replace", "delete"):
            ranges.append((i1, i2))
        elif tag == "insert" and i1 > 0:
            ranges.append((max(0, i1-1), i1))
    return ranges

# ── review doc 段落操作 ───────────────────────────────────────────────────────
def _add_separator(doc, article_id, title):
    sep = doc.add_paragraph()
    r = sep.add_run(f"═══ ARTICLE: {article_id} ║ {title} ═══")
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1A, 0x56, 0xAA)

def _copy_para(src_para, dest_doc):
    """複製段落 XML 到 dest_doc（保留原始格式）"""
    dest_doc.element.body.append(deepcopy(src_para._element))

def _add_highlight_and_hint(dest_doc, src_para, highlight_ranges, new_text):
    """
    複製 src_para 到 dest_doc，對 highlight_ranges 範圍的 runs 加黃螢光筆。
    下方追加一行灰色提示：→ 改為：{new_text}
    """
    _copy_para(src_para, dest_doc)
    new_para = dest_doc.paragraphs[-1]

    # 標記螢光筆（run 粒度）
    pos = 0
    for run in new_para.runs:
        rs, re_ = pos, pos + len(run.text)
        for hl_s, hl_e in highlight_ranges:
            if rs < hl_e and re_ > hl_s:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                break
        pos = re_

    # 提示行
    hint = dest_doc.add_paragraph()
    hr = hint.add_run(f"    → 改為：{new_text}")
    hr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    hr.font.size = Pt(9)
    hr.font.italic = True

# ── 在 Word 資料夾中找對應 .docx ──────────────────────────────────────────────
def find_word_file(folder, article_id, title=None):
    """
    Supabase 文章 ID 就是檔名（去掉 .docx），直接精確比對。
    例：article_id="7-3火燒島的美麗風景" → 7-3火燒島的美麗風景.docx
    """
    target = Path(folder) / f"{article_id}.docx"
    return target if target.exists() else None

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 步驟一：review
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def cmd_review(issue_number, word_folder=None):
    # 自動找資料夾
    if word_folder:
        folder = Path(word_folder)
    else:
        folder = auto_issue_folder(issue_number)
        if not folder:
            print(f"❌  找不到 G:/我的雲端硬碟/資料/無境界者/雜誌/{int(issue_number):02d}-* 資料夾", file=sys.stderr)
            sys.exit(1)
        print(f"📁  自動使用資料夾：{folder}")

    if not folder.is_dir():
        print(f"❌  資料夾不存在：{folder}", file=sys.stderr)
        sys.exit(1)

    print(f"⬇   從 Supabase 抓取第 {issue_number} 期文章…")
    articles = fetch_issue_articles(issue_number)
    print(f"    共 {len(articles)} 篇\n")

    review_doc = Document()
    # 移除預設空白段落
    for p in list(review_doc.paragraphs):
        p._element.getparent().remove(p._element)

    change_log     = []
    total_changes  = 0

    for art in articles:
        art_id    = art["id"]
        art_title = art.get("title", "（無標題）")

        wfile = find_word_file(folder, art_id, art_title)
        if not wfile:
            print(f"  ⚠️   [{art_id}] 找不到對應 Word：{art_title}")
            continue

        print(f"  📄  {art_title}")
        print(f"       檔案：{wfile.name}")

        try:
            word_doc = Document(str(wfile))
        except Exception as e:
            print(f"       ❌  無法開啟：{e}")
            continue

        word_entries = collect_nonempty_paras(word_doc)
        html_paras   = extract_html_paras(art.get("content", ""))
        diff_map     = align_paras(word_entries, html_paras)  # {word_orig_idx: chg}

        _add_separator(review_doc, art_id, art_title)

        art_changes = []

        # 按照 Word 文件原始段落順序加入 review doc
        for i, para in enumerate(word_doc.paragraphs):
            if i in diff_map:
                chg    = diff_map[i]
                w_text = chg["w_text"]
                h_text = chg["h_text"]
                sim    = chg["sim"]

                rv_idx = len(review_doc.paragraphs)
                ranges = changed_ranges_in_old(w_text, h_text)
                _add_highlight_and_hint(review_doc, para, ranges, h_text)

                art_changes.append({
                    "review_para_idx":    rv_idx,
                    "word_orig_para_idx": i,
                    "word_text":          w_text,
                    "supabase_text":      h_text,
                    "similarity":         round(sim, 3),
                })
                total_changes += 1
                print(f"       ✏️   段落 {i+1}（{sim:.0%}）")
                print(f"            舊：{w_text[:60]}{'…' if len(w_text)>60 else ''}")
                print(f"            新：{h_text[:60]}{'…' if len(h_text)>60 else ''}")
            else:
                _copy_para(para, review_doc)

        if art_changes:
            change_log.append({
                "article_id": art_id,
                "title":      art_title,
                "word_file":  str(wfile),
                "changes":    art_changes,
            })

    # 輸出
    out_dir  = folder.parent
    doc_out  = out_dir / f"review_issue{issue_number}.docx"
    json_out = out_dir / f"review_issue{issue_number}.json"

    review_doc.save(str(doc_out))
    with open(json_out, "w", encoding="utf-8") as f:
        json.dump(change_log, f, ensure_ascii=False, indent=2)

    print(f"\n{'─'*60}")
    print(f"✅  完成！共 {total_changes} 個改動已用黃螢光筆標示")
    print(f"    審稿檔：{doc_out}")
    print(f"    改動紀錄：{json_out}")
    print()
    print("下一步：")
    print("  1. 開啟審稿檔，確認黃螢光筆段落正確")
    print("  2. 若某段不應更改，在 Word 裡移除該段的黃螢光筆")
    print(f"  3. 儲存後執行：")
    print(f"     python scripts/patch_docx.py apply {doc_out.name} --inplace")

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 步驟二：apply
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def _has_yellow(para):
    return any(r.font.highlight_color == WD_COLOR_INDEX.YELLOW
               for r in para.runs)

def cmd_apply(review_doc_path, word_folder=None, inplace=False):
    review_doc_path = Path(review_doc_path)
    if not review_doc_path.exists():
        print(f"❌  找不到審稿檔：{review_doc_path}", file=sys.stderr)
        sys.exit(1)

    json_path = review_doc_path.with_suffix(".json")
    if not json_path.exists():
        print(f"❌  找不到改動紀錄：{json_path}", file=sys.stderr)
        sys.exit(1)

    with open(json_path, encoding="utf-8") as f:
        change_log = json.load(f)

    review_doc   = Document(str(review_doc_path))
    review_paras = review_doc.paragraphs

    print(f"📄  審稿檔：{review_doc_path.name}")
    print(f"    共 {len(change_log)} 篇文章有改動紀錄\n")

    applied = skipped = failed = 0
    flagged = []

    for art_log in change_log:
        art_id    = art_log["article_id"]
        art_title = art_log["title"]
        wfile     = Path(art_log["word_file"])

        # 若記錄的路徑不存在，嘗試在 word_folder 或 json 同層資料夾重新找
        if not wfile.exists():
            search_folder = Path(word_folder) if word_folder else json_path.parent
            found = find_word_file(search_folder, art_id)
            if found:
                wfile = found

        if not wfile.exists():
            print(f"  ⚠️   找不到原始 Word：{art_id} | {art_title}")
            continue

        print(f"  📝  {art_title}  ({art_id})")

        try:
            word_doc = Document(str(wfile))
        except Exception as e:
            print(f"       ❌  無法開啟：{e}")
            continue

        doc_modified = False

        for chg in art_log["changes"]:
            rv_idx      = chg["review_para_idx"]
            w_text_orig = chg["word_text"]
            s_text_new  = chg["supabase_text"]

            # 確認螢光筆是否仍在
            if rv_idx >= len(review_paras):
                print(f"       ⚠️   review 索引超範圍（{rv_idx}），跳過")
                skipped += 1
                continue

            if not _has_yellow(review_paras[rv_idx]):
                print(f"       ─   螢光筆已移除，跳過")
                print(f"            {w_text_orig[:55]}")
                skipped += 1
                continue

            # 在原始 Word 中找對應段落（精確比對優先）
            target_para = None
            best_sim = 0.0
            for para in word_doc.paragraphs:
                pt = para_text(para)
                if pt == w_text_orig:
                    target_para = para
                    best_sim = 1.0
                    break
                sim = difflib.SequenceMatcher(None, pt, w_text_orig).ratio()
                if sim > best_sim:
                    best_sim, target_para = sim, para

            if target_para is None or best_sim < 0.8:
                reason = f"找不到對應段落（最高相似度 {best_sim:.0%}）"
                print(f"       ⚠️   {reason}")
                print(f"            尋找：{w_text_orig[:55]}")
                flagged.append((art_id, w_text_orig, s_text_new, reason))
                failed += 1
                continue

            ok, msg = try_patch_para(target_para, s_text_new)
            if ok:
                doc_modified = True
                applied += 1
                print(f"       ✅  {msg}")
                print(f"            {w_text_orig[:45]} → {s_text_new[:45]}")
            else:
                flagged.append((art_id, w_text_orig, s_text_new, msg))
                failed += 1
                print(f"       ⚠️   {msg}")

        if doc_modified:
            if inplace:
                word_doc.save(str(wfile))
                print(f"       💾  已覆蓋：{wfile.name}")
            else:
                out_path = wfile.parent / (wfile.stem + "_patched" + wfile.suffix)
                word_doc.save(str(out_path))
                print(f"       💾  輸出：{out_path.name}")

    print(f"\n{'─'*60}")
    print(f"套用 {applied} 個，使用者拒絕 {skipped} 個，失敗 {failed} 個")

    if flagged:
        print(f"\n⚠️   以下需要手動處理：")
        for art_id, old_t, new_t, reason in flagged:
            print(f"  [{art_id}] {reason}")
            print(f"    舊：{old_t[:60]}")
            print(f"    新：{new_t[:60]}")

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 入口
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def main():
    args = sys.argv[1:]
    if not args or args[0] in ("-h", "--help"):
        print(__doc__)
        return

    cmd = args[0].lower()

    if cmd == "review":
        if len(args) < 2:
            print("用法：python scripts/patch_docx.py review <期數> [Word資料夾]")
            sys.exit(1)
        folder = args[2] if len(args) >= 3 else None
        cmd_review(args[1], folder)

    elif cmd == "apply":
        if len(args) < 2:
            print("用法：python scripts/patch_docx.py apply <review_doc.docx> [Word資料夾] [--inplace]")
            sys.exit(1)
        inplace = "--inplace" in args
        folder  = next((a for a in args[2:] if not a.startswith("--")), None)
        cmd_apply(args[1], folder, inplace=inplace)

    else:
        print(f"❌  未知指令：{cmd}（請用 review 或 apply）")
        sys.exit(1)

if __name__ == "__main__":
    main()
