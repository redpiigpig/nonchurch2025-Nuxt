# -*- coding: utf-8 -*-
"""
render_meta_docx.py — 從範例 docx template 渲染當期投稿資訊／編輯資訊

呼叫：
    python render_meta_docx.py <input.json> <output.docx>

input.json 結構：
{
  "type": "submission_info" | "editorial_info",
  "issue": {
    "id": 8,
    "title": "地上神國與人間佛教",
    "date": "2026年03-04月號",
    "cfp_title": "...",
    "cfp_image": "https://res.cloudinary.com/.../xxx.png",
    "cfp_theme": "段落1\n\n段落2\n\n段落3",
    "cfp_deadline": "2026年6月15日"
  },
  "finance": {                      # 僅 editorial_info 需要；可為 null
    "dateRange": "115.02.16 – 115.04.15",
    "rows": [
      { "id": "...", "date": "...", "type": "...", "item": "...",
        "category": "...", "unitPrice": "...", "qty": "...",
        "total": 1234, "balance": -5678 },
      ...
    ]
  }
}
"""

import io
import json
import sys
import urllib.request
from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

TEMPLATE_DIR = None  # 由呼叫方注入；fallback 取相對於本檔案的 ../templates

NS_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


# ─────────────────────────────────────────────────────────────────
# 共用 helpers
# ─────────────────────────────────────────────────────────────────

def fmt_num(n):
    if n is None or n == "":
        return ""
    try:
        return "{:,}".format(int(n))
    except (TypeError, ValueError):
        return str(n)


def first_run_rPr_clone(p_element):
    """取段落第一個 run 的 rPr 拷貝（用作新 run 的樣式模板）。沒有就回 None。"""
    r = p_element.find(qn("w:r"))
    if r is None:
        return None
    rPr = r.find(qn("w:rPr"))
    return deepcopy(rPr) if rPr is not None else None


def build_run(text, rPr_clone=None):
    """組一個 <w:r> element，含 rPr 與 w:t（text 保留前後空白）。"""
    r = OxmlElement("w:r")
    if rPr_clone is not None:
        r.append(rPr_clone)
    t = OxmlElement("w:t")
    t.text = text or ""
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    return r


def replace_full_paragraph(paragraph, new_text):
    """段落整段同樣式時用：保留段落 pPr 與第一個 run 的 rPr，寫入 new_text。"""
    p_el = paragraph._p
    rPr_clone = first_run_rPr_clone(p_el)
    # 移除所有現有的 r、hyperlink、smartTag
    for child in list(p_el):
        if child.tag in (qn("w:r"), qn("w:hyperlink"), qn("w:smartTag")):
            p_el.remove(child)
    new_r = build_run(new_text, rPr_clone)
    p_el.append(new_r)


def replace_after_prefix_run(paragraph, new_suffix):
    """段落形如 [run0:粗體前綴][run1..n:非粗體內文]，保留 run0，把後面整段替換為 new_suffix。

    用第二個 run（如有）的 rPr 當樣式；沒有第二個 run 就新建一個，rPr 為 None。
    """
    p_el = paragraph._p
    runs = p_el.findall(qn("w:r"))
    if not runs:
        # 沒 run 就退化成 full replace（其實不該發生）
        replace_full_paragraph(paragraph, new_suffix)
        return
    if len(runs) == 1:
        # 只有一個 run，這 run 已含全文；當作 full replace
        replace_full_paragraph(paragraph, new_suffix)
        return
    suffix_rPr = runs[1].find(qn("w:rPr"))
    suffix_rPr_clone = deepcopy(suffix_rPr) if suffix_rPr is not None else None
    # 移除 runs[1:]（連 hyperlink 之類也清）
    keep = runs[0]
    for child in list(p_el):
        if child is keep:
            continue
        if child.tag in (qn("w:r"), qn("w:hyperlink"), qn("w:smartTag")):
            p_el.remove(child)
    # 在 keep 後插入新 run
    new_r = build_run(new_suffix, suffix_rPr_clone)
    keep.addnext(new_r)


def set_cell_text(tc_element, text):
    """設儲存格純文字，保留第一個 paragraph 的 pPr 與第一個 run 的 rPr。"""
    paragraphs = tc_element.findall(qn("w:p"))
    if not paragraphs:
        # 建立一個段
        p = OxmlElement("w:p")
        r = build_run(text)
        p.append(r)
        tc_element.append(p)
        return
    keep_p = paragraphs[0]
    # 刪除其他 paragraph
    for p in paragraphs[1:]:
        tc_element.remove(p)
    # 取第一個 run rPr 樣式
    rPr_clone = first_run_rPr_clone(keep_p)
    # 移除原 runs
    for child in list(keep_p):
        if child.tag in (qn("w:r"), qn("w:hyperlink"), qn("w:smartTag")):
            keep_p.remove(child)
    keep_p.append(build_run(text, rPr_clone))


def replace_image_blob(doc, target_filename_suffix, new_bytes):
    """把 doc 中 target_ref 結尾為 target_filename_suffix 的 image part bytes 替換。"""
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype and rel.target_ref.endswith(target_filename_suffix):
            rel.target_part._blob = new_bytes
            return True
    return False


def http_download(url):
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    return urllib.request.urlopen(req, timeout=20).read()


def split_theme_paragraphs(theme_text):
    """把 cfp_theme 多行字串拆成段落清單（依空行或單一換行分段）。"""
    if not theme_text:
        return []
    # 先依雙換行（段落界線），再對每段去除全形空白前綴
    raw_segments = [seg.strip() for seg in theme_text.split("\n\n")]
    if len(raw_segments) == 1:
        # fallback：依單一換行
        raw_segments = [seg.strip() for seg in theme_text.split("\n")]
    return [s for s in raw_segments if s]


# ─────────────────────────────────────────────────────────────────
# 頁首 patching：頁首的文字方塊（textbox）內藏 <w:p>，
# python-docx 的 section.header.paragraphs 抓不到，因此直接走 element.iter。
# ─────────────────────────────────────────────────────────────────

def _full_text_of(p_element):
    return "".join((t.text or "") for t in p_element.iter(qn("w:t")))


def _rebuild_paragraph_text(p_element, new_text):
    """整段重寫：保留 pPr 與第一個 run 的 rPr，刪除其他 runs 與 hyperlink。"""
    rPr_clone = first_run_rPr_clone(p_element)
    for child in list(p_element):
        if child.tag in (qn("w:r"), qn("w:hyperlink"), qn("w:smartTag")):
            p_element.remove(child)
    p_element.append(build_run(new_text, rPr_clone))


def patch_header_paragraphs(hdr_element, anchor_substring, new_text):
    """對 hdr 內所有「leaf paragraph」（不含巢狀 <w:p>）做替換。

    跳過容器段（textbox 外層的 <w:p>），避免把整個 textbox 結構連同子段一起吃掉。
    """
    count = 0
    for p in list(hdr_element.iter(qn("w:p"))):
        # 容器段（如 textbox 外層）內部還有子 <w:p>，跳過
        if p.find(".//" + qn("w:p")) is not None:
            continue
        full = _full_text_of(p)
        if anchor_substring in full:
            _rebuild_paragraph_text(p, new_text)
            count += 1
    return count


def format_header_date(raw):
    """頁首日期：'2026年01-02月號' / '2026年02月號' / '2026.04' → 'YYYY.MM-MM' 或 'YYYY.MM'。"""
    if not raw:
        return ""
    import re
    s = str(raw)
    m = re.search(r"(\d{4})年(\d{1,2})-(\d{1,2})月", s)
    if m:
        return "{}.{:02d}-{:02d}".format(m.group(1), int(m.group(2)), int(m.group(3)))
    m = re.search(r"(\d{4})年(\d{1,2})月", s)
    if m:
        return "{}.{:02d}".format(m.group(1), int(m.group(2)))
    m = re.search(r"(\d{4})\.(\d{1,2})-(\d{1,2})", s)
    if m:
        return "{}.{:02d}-{:02d}".format(m.group(1), int(m.group(2)), int(m.group(3)))
    m = re.search(r"(\d{4})\.(\d{1,2})", s)
    if m:
        # 雙月刊推算：3 月刊→3-4 月、4 月刊→3-4 月、5 月刊→5-6 月、6 月刊→5-6 月
        mo = int(m.group(2))
        lo = mo - 1 if mo % 2 == 0 else mo
        hi = mo if mo % 2 == 0 else mo + 1
        return "{}.{:02d}-{:02d}".format(m.group(1), lo, hi)
    return s


def extract_seq_from_id(article_id):
    """'8-17投稿資訊' → '17'"""
    import re
    m = re.match(r"\d+-(\d+)", str(article_id or ""))
    return m.group(1) if m else ""


def patch_meta_headers(doc, issue, article_id, label):
    """更新所有 section 的頁首：
        - 「無境界者｜Vol. 7（2026.01-02）」段 → 套當期 issue id 與日期
        - 「火燒島上的《耶穌傳》」段 → 套 issue.title
        - header2 的「17投稿資訊」段 → 套 article_id 序號 + label
    label 例：'投稿資訊' 或 '編輯資訊'。
    """
    issue_id = str(issue.get("id") or "")
    issue_title = issue.get("title") or ""
    date_label = format_header_date(issue.get("date") or "")
    seq = extract_seq_from_id(article_id) or ""

    masthead = "無境界者｜Vol. {}（{}）".format(issue_id, date_label) if issue_id else ""
    seq_label = "{}{}".format(seq, label) if seq else label

    for section in doc.sections:
        for hdr_attr in ("header", "first_page_header", "even_page_header"):
            hdr = getattr(section, hdr_attr, None)
            if hdr is None:
                continue
            hdr_el = hdr._element
            if masthead:
                patch_header_paragraphs(hdr_el, "Vol.", masthead)
            if issue_title:
                patch_header_paragraphs(hdr_el, "火燒島上的", issue_title)
            if seq_label:
                patch_header_paragraphs(hdr_el, label, seq_label)


# ─────────────────────────────────────────────────────────────────
# 投稿資訊
# ─────────────────────────────────────────────────────────────────

def render_submission_info(template_path, issue, article_id, output_path):
    doc = Document(template_path)
    patch_meta_headers(doc, issue, article_id, "投稿資訊")

    # p[7] 「下期主題：「地上神國與人間佛教」」
    cfp_title = (issue.get("cfp_title") or "（待編輯）").strip()
    replace_full_paragraph(doc.paragraphs[7], "下期主題：「{}」".format(cfp_title))

    # p[8] cfp_image: 替換 image1.png blob
    cfp_image = issue.get("cfp_image")
    if cfp_image:
        try:
            img_bytes = http_download(cfp_image)
            replace_image_blob(doc, "image1.png", img_bytes)
        except Exception as e:
            print("Warning: cfp_image download failed: {}".format(e), file=sys.stderr)

    # p[9..12] cfp_theme 多段：保留 p[9] 作模板，刪除 p[10..12]，插入新內容
    theme_paragraphs = split_theme_paragraphs(issue.get("cfp_theme") or "")
    if not theme_paragraphs:
        theme_paragraphs = ["（請於期刊主題管理填入下期徵稿說明，再回此頁同步。）"]

    template_p_el = doc.paragraphs[9]._p
    template_pPr = template_p_el.find(qn("w:pPr"))
    template_rPr = first_run_rPr_clone(template_p_el)

    # 收集 p[10..12] 的 element 待刪除
    to_delete = [doc.paragraphs[i]._p for i in (10, 11, 12)]

    # 重設 p[9] 為 theme_paragraphs[0]
    replace_full_paragraph(doc.paragraphs[9], theme_paragraphs[0])

    # 對 theme_paragraphs[1:]：clone p[9] 結構並插入
    insert_after = doc.paragraphs[9]._p
    for txt in theme_paragraphs[1:]:
        new_p = OxmlElement("w:p")
        if template_pPr is not None:
            new_p.append(deepcopy(template_pPr))
        new_p.append(build_run(txt, deepcopy(template_rPr) if template_rPr is not None else None))
        insert_after.addnext(new_p)
        insert_after = new_p

    # 刪除原本 p[10..12]
    for el in to_delete:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

    # 截稿期限段：插入後的新段落會把 index 推移，重新尋找含「📌截稿期限」的段落
    deadline = (issue.get("cfp_deadline") or "（待編輯）").strip()
    for p in doc.paragraphs:
        if p.text.startswith("📌截稿期限") or p.text.startswith("📌 截稿期限"):
            replace_full_paragraph(p, "📌截稿期限：{}".format(deadline))
            break

    doc.save(output_path)


# ─────────────────────────────────────────────────────────────────
# 編輯資訊
# ─────────────────────────────────────────────────────────────────

ZH_NUM = ["", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]


def to_zh_issue_num(idv):
    try:
        n = int(idv)
    except (TypeError, ValueError):
        return str(idv)
    if 1 <= n <= 10:
        return ZH_NUM[n]
    if 11 <= n < 20:
        return "十" + ZH_NUM[n - 10]
    return str(n)


def format_publish_date(raw):
    """'2026年03-04月號' / '2026.04' → '2026年3月'"""
    if not raw:
        return ""
    import re
    m = re.search(r"(\d{4})年(\d{1,2})", raw)
    if m:
        return "{}年{}月".format(m.group(1), int(m.group(2)))
    m = re.search(r"(\d{4})\.(\d{1,2})", raw)
    if m:
        return "{}年{}月".format(m.group(1), int(m.group(2)))
    return raw


def find_finance_table(doc):
    """財務表辨識：含 9 欄、且 row[0] 第一格文字為「編號」。"""
    for t in doc.tables:
        try:
            if len(t.columns) >= 9 and t.rows[0].cells[0].text.strip().startswith("編號"):
                return t
        except Exception:
            continue
    return None


def render_editorial_info(template_path, issue, finance, article_id, output_path):
    doc = Document(template_path)
    patch_meta_headers(doc, issue, article_id, "編輯資訊")

    # p[12] 「（114.12.15-115.02.15）」 → 財務徵信日期範圍
    if finance and finance.get("dateRange"):
        for p in doc.paragraphs:
            txt = p.text.strip()
            if txt.startswith("（") and ("." in txt or "-" in txt) and "）" in txt:
                # 經驗辨識：（YYY.MM.DD-YYY.MM.DD）
                if any(c.isdigit() for c in txt[:10]):
                    replace_full_paragraph(p, "（{}）".format(finance["dateRange"]))
                    break

    # p[19] 卷期：保留粗體前綴「卷　期：」
    issue_id = issue.get("id")
    issue_title = issue.get("title") or ""
    suffix_volume = "第{}期：{}".format(to_zh_issue_num(issue_id), issue_title)
    for p in doc.paragraphs:
        if p.text.startswith("卷　期：") or p.text.startswith("卷 期："):
            replace_after_prefix_run(p, suffix_volume)
            break

    # p[20] 出版日期：保留粗體前綴
    pub_date = format_publish_date(issue.get("date"))
    for p in doc.paragraphs:
        if p.text.startswith("出版日期："):
            replace_after_prefix_run(p, pub_date)
            break

    # p[26] Copyright 年份替換
    year = pub_date[:4] if pub_date and pub_date[:4].isdigit() else ""
    if year:
        for p in doc.paragraphs:
            if p.text.startswith("Copyright"):
                replace_full_paragraph(
                    p,
                    "Copyright © {} Faith Without Borders. All rights reserved.".format(year),
                )
                break

    # 財務表：保留 row[0]（header）、用 row[1] 作模板重建內容
    # 不論 finance 是否為 null，都清空 template 內示範用的舊 body rows，
    # 避免新期次 docx 還殘留第 7 期的範例資料。
    table = find_finance_table(doc)
    if table is not None and len(table.rows) >= 2:
        tbl_el = table._tbl
        template_tr = deepcopy(table.rows[1]._tr)
        # 刪除舊的 body rows（rows[1:]）
        for row in list(table.rows)[1:]:
            tbl_el.remove(row._tr)
        # 插入新 rows（finance 為空時 docx 表只剩 header）
        finance_rows = (finance or {}).get("rows") or []
        for row_data in finance_rows:
            # 編號欄用 display_seq（client 計算後傳入），fallback 才用 id
            seq = row_data.get("display_seq") or row_data.get("id") or ""
            values = [
                str(seq),
                str(row_data.get("date") or ""),
                str(row_data.get("type") or ""),
                str(row_data.get("item") or ""),
                str(row_data.get("category") or ""),
                str(row_data.get("unitPrice") or ""),
                str(row_data.get("qty") or ""),
                fmt_num(row_data.get("total")),
                fmt_num(row_data.get("balance")),
            ]
            new_tr = deepcopy(template_tr)
            tcs = new_tr.findall(qn("w:tc"))
            for tc, val in zip(tcs, values):
                set_cell_text(tc, val)
            tbl_el.append(new_tr)

    doc.save(output_path)


# ─────────────────────────────────────────────────────────────────
# entry
# ─────────────────────────────────────────────────────────────────

def main():
    if len(sys.argv) < 3:
        print("Usage: python render_meta_docx.py <input.json> <output.docx>", file=sys.stderr)
        sys.exit(2)

    input_path, output_path = sys.argv[1], sys.argv[2]
    with open(input_path, "r", encoding="utf-8") as f:
        payload = json.load(f)

    meta_type = payload.get("type")
    issue = payload.get("issue") or {}
    finance = payload.get("finance")
    article_id = payload.get("id") or ""

    import os
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    if meta_type == "submission_info":
        tpl = os.path.join(base_dir, "templates", "submission_info_template.docx")
        render_submission_info(tpl, issue, article_id, output_path)
    elif meta_type == "editorial_info":
        tpl = os.path.join(base_dir, "templates", "editorial_info_template.docx")
        render_editorial_info(tpl, issue, finance, article_id, output_path)
    else:
        print("Unsupported type: {}".format(meta_type), file=sys.stderr)
        sys.exit(2)

    print(output_path)


if __name__ == "__main__":
    main()
