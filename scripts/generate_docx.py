"""
無境界者雜誌 - 專業 Word 排版生成器
完全符合 form.md 格式規範

━━ 函數索引（中文對照）━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

【頁面初始化】
  __init__ / _setup_page / _setup_styles      L.28   初始化、頁面設定、樣式
  _inject_footnote_ref_style                  L.66   腳注引用上標樣式
  _clear_numbering                            L.88   清除列表自動編號

【基礎工具】
  _add_blank_line                             L.102  插入空白行
  _apply_font(run, ascii, east, size, bold)   L.116  套用字型（英文/中文/大小/粗體）
  _hex_rgb                                    L.140  十六進位色碼轉 RGBColor
  _download_image(src)                        L.347  下載 Cloudinary 圖片 → BytesIO
  _download_and_crop_square(src)              L.362  下載並裁切正方形圖片
  _split_special_chars(text)                  L.1468 拆分含 ‧／∕ 的文字（新細明體）
  _split_emoji(text)                          L.1478 拆分 emoji（Segoe UI Symbol）
  _should_skip_blank_before                   L.1320 判斷是否略過前置空行

【腳注系統】
  set_footnotes(footnotes)                    L.146  傳入腳注 JSONB 陣列
  _add_footnote_ref(paragraph, fn_num_str)    L.149  在段落插入 Word 腳注引用編號
  _fn_runs_xml(text)                          L.166  腳注內文 inline XML 生成
  _finalize_footnotes                         L.236  完成所有腳注寫入

【文章結構（依輸出順序）】
  add_category_tag(category, color_hex)       L.575  文章分類標籤（彩色色塊）
  add_title(title)                            L.585  主標題
  add_subtitle(subtitle)                      L.593  副標題
  add_decoration_line                         L.606  裝飾橫線
  add_author(author, author_title, remark)    L.625  作者、頭銜、備註（備註支援腳注）
  add_keywords(keywords)                      L.663  關鍵字行
  add_content(content)                        L.724  主要內文（呼叫下方各處理器）
  add_header_footer(...)                      L.1608 頁首頁尾（期數、標題、頁碼）

【內文解析器】
  _split_into_segments(content)               L.679  將 HTML 切分為段落類型清單
  _find_close_tag(content, from_pos, tag)     L.704  找配對的閉合標籤
  _dispatch_div(html)                         L.767  依 class 派發 <div> 到對應處理器
  _add_paragraph_element(html)                L.808  處理 <p> 段落（含首行縮排邏輯）
  _process_line(line)                         L.1347 處理純文字行（含縮排判斷）
  _add_inline(p, text, east, ascii)           L.1498 解析 inline 標記（粗體/楷體/斜體/腳注）

【自訂區塊】
  _add_custom_divider                         L.847  自訂分隔線（── ）
  _add_book_quote(html)                       L.869  書本引言（藍色左線、楷書體）
  _add_book_box(html)                         L.927  書籍簡介框（2欄：文字2/3 + 封面圖1/3）
  _add_reference_box(html)                    L.1037 參考資料框（綠色左線）
  _add_theme_image(html)                      L.1082 主題圖片（置中大圖）
  _add_author_profile(html)                   L.1097 作者簡介（頭像 + 文字）
  _add_right_aligned(text)                    L.1210 置右文字段落
  _add_blockquote(text, rel_text)             L.1447 一般引言（標楷體，rel 置右）

【圖片系統】
  _add_figure(html)                           L.384  圖片區塊（含圖說）
  _add_figure_textbox(...)                    L.450  圖片文字框（浮動排版）
  _add_caption_runs(paragraph, lines)         L.563  圖說文字渲染
  _insert_float_image(p, stream, width, dir)  L.295  插入浮動圖片（繞排）

【表格】
  _add_table(html)                            L.1230 HTML <table> → Word 表格
  _add_section_title(inner_html)              L.1385 <h3> 小標題（14pt 粗體，支援腳注）
  _add_bullet_line(text)                      L.1335 項目清單行

【頁面輔助】
  _add_portrait_double_border(run)            L.1256 雙框線效果
  _add_picture_border(run, border_pt, color)  L.1292 圖片框線
  _insert_page_number(paragraph, size)        L.1700 插入頁碼欄位

【進入點】
  generate_article_docx(article_data, path)   L.1725 主函數（組合所有步驟）

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys
import json
import re
import html as _html_mod
import urllib.request
import io
from PIL import Image as PILImage

# 強制 stdout/stderr 使用 UTF-8，避免 Windows CP950 無法輸出 emoji
if sys.stdout.encoding != 'utf-8':
    sys.stdout = open(sys.stdout.fileno(), mode='w', encoding='utf-8', buffering=1)
if sys.stderr.encoding != 'utf-8':
    sys.stderr = open(sys.stderr.fileno(), mode='w', encoding='utf-8', buffering=1)


class ProfessionalDocxGenerator:

    def __init__(self):
        self.doc = Document()
        self._setup_page()
        self._setup_styles()
        self._inject_footnote_ref_style()
        self._clear_numbering()  # 清除預設列表定義，防止 keepNext 觸發圓點
        self._footnotes = []    # [(word_id, text), ...]
        self._fn_counter = 1
        self._fn_map = {}       # "[^N]" → text
        self._img_counter = 1

    # ── 頁面設定 ────────────────────────────────────────────

    def _setup_page(self):
        section = self.doc.sections[0]
        section.page_width  = Cm(18.2)
        section.page_height = Cm(25.7)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)
        # 頁首距頁頂 0.5cm，頁尾距頁底 1cm
        section.header_distance = Cm(0.5)
        section.footer_distance = Cm(1.0)

    def _setup_styles(self):
        """正文基礎樣式：NSimSun + Times New Roman，1.5 倍行距"""
        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        rPr = style.element.get_or_add_rPr()
        rPr.get_or_add_rFonts().set(qn('w:eastAsia'), 'NSimSun')
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)

    def _inject_footnote_ref_style(self):
        """注入 FootnoteReference 字元樣式（上標），python-docx 空白文件不含此樣式"""
        styles_el = self.doc.styles.element
        for style in styles_el.findall(qn('w:style')):
            if style.get(qn('w:styleId')) == 'FootnoteReference':
                return  # 已存在，不重複注入
        fn_style = OxmlElement('w:style')
        fn_style.set(qn('w:type'), 'character')
        fn_style.set(qn('w:styleId'), 'FootnoteReference')
        name_el = OxmlElement('w:name')
        name_el.set(qn('w:val'), 'footnote reference')
        fn_style.append(name_el)
        based = OxmlElement('w:basedOn')
        based.set(qn('w:val'), 'DefaultParagraphFont')
        fn_style.append(based)
        rPr = OxmlElement('w:rPr')
        vertAlign = OxmlElement('w:vertAlign')
        vertAlign.set(qn('w:val'), 'superscript')
        rPr.append(vertAlign)
        fn_style.append(rPr)
        styles_el.append(fn_style)

    def _clear_numbering(self):
        """清除文件預設的所有列表編號定義（abstractNum / num）
        防止 Word 看到 keepNext 時自動套用列表圓點樣式"""
        try:
            num_part = self.doc.part.numbering_part
            if num_part is None:
                return
            el = num_part._element
            for tag in (qn('w:abstractNum'), qn('w:num')):
                for child in list(el.findall(tag)):
                    el.remove(child)
        except AttributeError:
            pass  # 文件本來就沒有 numbering part，不需要處理

    def _add_blank_line(self):
        """插入一個空行，明確設定 1.5 倍行距（line=360 auto）"""
        p = self.doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:before'), '0')
        sp.set(qn('w:after'),  '0')
        sp.set(qn('w:line'),   '360')    # 240 × 1.5 = 360
        sp.set(qn('w:lineRule'), 'auto')
        pPr.append(sp)
        return p

    # ── 字型輔助 ────────────────────────────────────────────

    def _apply_font(self, run, ascii_font, east_font=None, size=12,
                    bold=False, italic=False, color=None, superscript=False):
        run.font.name  = ascii_font
        run.font.size  = Pt(size)
        run.font.bold  = bold
        run.font.italic = italic
        if superscript:
            run.font.superscript = True
        rPr = run._element.get_or_add_rPr()
        rf = rPr.get_or_add_rFonts()
        rf.set(qn('w:ascii'), ascii_font)
        rf.set(qn('w:hAnsi'), ascii_font)
        if east_font:
            rf.set(qn('w:eastAsia'), east_font)
        if color:
            run.font.color.rgb = RGBColor(*color)
        # 明確設定 w:szCs（複合腳本字型大小），確保中文字也套用正確字級
        half_pts = str(int(size * 2))
        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = OxmlElement('w:szCs')
            rPr.append(szCs)
        szCs.set(qn('w:val'), half_pts)

    def _hex_rgb(self, hex_color):
        h = hex_color.lstrip('#')
        return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

    # ── 腳注 ────────────────────────────────────────────────

    def set_footnotes(self, footnotes):
        self._fn_map = {str(fn.get('id', '')): fn.get('text', '') for fn in (footnotes or [])}

    def _add_footnote_ref(self, paragraph, fn_num_str):
        """在段落插入 Word 腳注引用標記（右上角小字）"""
        fn_text  = self._fn_map.get(fn_num_str, f'[腳注 {fn_num_str}]')
        word_id  = self._fn_counter
        self._fn_counter += 1
        self._footnotes.append((word_id, fn_text))

        run = paragraph.add_run()
        rPr = OxmlElement('w:rPr')
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'FootnoteReference')
        rPr.append(rStyle)
        run._element.insert(0, rPr)
        fnRef = OxmlElement('w:footnoteReference')
        fnRef.set(qn('w:id'), str(word_id))
        run._element.append(fnRef)

    def _fn_runs_xml(self, text):
        """腳注文字轉 Word XML runs。
        支援 **bold** *kaiti* <b> <i>；<a href="..."> 剝除標籤只保留文字；其他 HTML 略過。"""
        def esc(s):
            return (str(s).replace('&','&amp;').replace('<','&lt;')
                          .replace('>','&gt;').replace('"','&quot;'))

        # 剝離超連結 <a ...>text</a> → text
        text = re.sub(r'<a\b[^>]*>(.*?)</a>', r'\1',
                      text, flags=re.IGNORECASE | re.DOTALL)

        BASE_RPR = (
            '<w:rPr><w:sz w:val="20"/><w:szCs w:val="20"/>'
            '<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"'
            ' w:eastAsia="PMingLiU"/></w:rPr>'
        )

        TOKEN = re.compile(
            r'(\*\*[^*\n]+?\*\*'
            r'|\*[^*\n]+?\*'
            r'|<(?:b|strong)>[^<]*</(?:b|strong)>'
            r'|<em>[^<]*</em>'
            r'|<i>[^<]*</i>'
            r'|<[^>]+>)',
            re.IGNORECASE)

        parts = []
        for seg in TOKEN.split(text):
            if not seg:
                continue
            bold_md     = re.fullmatch(r'\*\*([^*]+)\*\*', seg)
            kaiti_md    = re.fullmatch(r'\*([^*]+)\*', seg)
            bold_html   = re.fullmatch(r'<(?:b|strong)>([^<]*)</(?:b|strong)>', seg, re.IGNORECASE)
            em_html     = re.fullmatch(r'<em>([^<]*)</em>', seg, re.IGNORECASE)
            italic_html = re.fullmatch(r'<i>([^<]*)</i>', seg, re.IGNORECASE)
            html_tag    = re.fullmatch(r'<[^>]+>', seg)

            if bold_md or bold_html:
                c = (bold_md or bold_html).group(1)
                parts.append(
                    f'<w:r><w:rPr><w:b/><w:sz w:val="20"/><w:szCs w:val="20"/>'
                    f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"'
                    f' w:eastAsia="PMingLiU"/></w:rPr>'
                    f'<w:t xml:space="preserve">{esc(c)}</w:t></w:r>'
                )
            elif kaiti_md or em_html:
                c = (kaiti_md or em_html).group(1)
                parts.append(
                    f'<w:r><w:rPr><w:sz w:val="20"/><w:szCs w:val="20"/>'
                    f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"'
                    f' w:eastAsia="標楷體"/></w:rPr>'
                    f'<w:t xml:space="preserve">{esc(c)}</w:t></w:r>'
                )
            elif italic_html:
                c = italic_html.group(1)
                parts.append(
                    f'<w:r><w:rPr><w:i/><w:sz w:val="20"/><w:szCs w:val="20"/>'
                    f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"'
                    f' w:eastAsia="PMingLiU"/></w:rPr>'
                    f'<w:t xml:space="preserve">{esc(c)}</w:t></w:r>'
                )
            elif html_tag:
                pass  # 略過其他 HTML 標籤
            else:
                parts.append(
                    f'<w:r>{BASE_RPR}'
                    f'<w:t xml:space="preserve">{esc(seg)}</w:t></w:r>'
                )
        return ''.join(parts)

    def _finalize_footnotes(self):
        """寫入 word/footnotes.xml（Word 原生腳注）"""
        if not self._footnotes:
            return

        # 1字元凸排 = 10pt = 200 twips；行距單倍 = line=240 auto
        lines = [
            "<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\n",
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">\n',
            '<w:footnote w:type="separator" w:id="-1">'
            '<w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
            '<w:r><w:separator/></w:r></w:p></w:footnote>\n',
            '<w:footnote w:type="continuationSeparator" w:id="0">'
            '<w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
            '<w:r><w:continuationSeparator/></w:r></w:p></w:footnote>\n',
        ]
        for fn_id, fn_text in self._footnotes:
            # 前置空格 run（腳注編號後的空格）
            space_run = (
                '<w:r><w:rPr><w:sz w:val="20"/><w:szCs w:val="20"/>'
                '<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"'
                ' w:eastAsia="PMingLiU"/></w:rPr>'
                '<w:t xml:space="preserve"> </w:t></w:r>'
            )
            lines.append(
                f'<w:footnote w:id="{fn_id}">'
                f'<w:p>'
                f'<w:pPr>'
                f'<w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>'
                f'<w:ind w:left="200" w:hanging="200"/>'
                f'</w:pPr>'
                f'<w:r>'
                f'<w:rPr>'
                f'<w:rStyle w:val="FootnoteReference"/>'
                f'<w:sz w:val="20"/><w:szCs w:val="20"/>'
                f'</w:rPr>'
                f'<w:footnoteRef/>'
                f'</w:r>'
                + space_run
                + self._fn_runs_xml(fn_text) +
                f'</w:p>'
                f'</w:footnote>\n'
            )
        lines.append('</w:footnotes>')

        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        doc_part = self.doc.part
        fn_part = Part(
            PackURI('/word/footnotes.xml'),
            'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml',
            ''.join(lines).encode('utf-8'),
            doc_part.package,
        )
        FOOTNOTES_REL = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
        doc_part.relate_to(fn_part, FOOTNOTES_REL)

    # ── 浮動圖片 ────────────────────────────────────────────

    def _insert_float_image(self, paragraph, img_stream, width_cm, float_dir='right'):
        WP = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
        A  = 'http://schemas.openxmlformats.org/drawingml/2006/main'

        run = paragraph.add_run()
        run.add_picture(img_stream, width=Cm(width_cm))

        drawing = run._r.find(qn('w:drawing'))
        if drawing is None:
            return
        inline = drawing.find(f'{{{WP}}}inline')
        if inline is None:
            return

        extent      = inline.find(f'{{{WP}}}extent')
        docPr       = inline.find(f'{{{WP}}}docPr')
        cNvGrfFrmPr = inline.find(f'{{{WP}}}cNvGraphicFramePr')
        graphic     = inline.find(f'{{{A}}}graphic')

        cx = extent.get('cx') if extent is not None else str(int(Cm(width_cm)))
        cy = extent.get('cy') if extent is not None else str(int(Cm(width_cm) * 1.4))

        img_id = self._img_counter; self._img_counter += 1
        if docPr is not None:
            docPr.set('id', str(img_id))
            docPr.set('name', f'Picture {img_id}')

        anchor = OxmlElement('wp:anchor')
        for k, v in [('distT','114300'),('distB','114300'),('distL','114300'),
                     ('distR','114300'),('simplePos','0'),('relativeHeight','251658240'),
                     ('behindDoc','0'),('locked','0'),('layoutInCell','1'),('allowOverlap','1')]:
            anchor.set(k, v)

        sp = OxmlElement('wp:simplePos'); sp.set('x','0'); sp.set('y','0'); anchor.append(sp)
        posH = OxmlElement('wp:positionH'); posH.set('relativeFrom','margin')
        al = OxmlElement('wp:align'); al.text = float_dir; posH.append(al); anchor.append(posH)
        posV = OxmlElement('wp:positionV'); posV.set('relativeFrom','paragraph')
        po = OxmlElement('wp:posOffset'); po.text = '0'; posV.append(po); anchor.append(posV)
        ext = OxmlElement('wp:extent'); ext.set('cx', cx); ext.set('cy', cy); anchor.append(ext)
        ef = OxmlElement('wp:effectExtent')
        for a in ('l','t','r','b'): ef.set(a,'0')
        anchor.append(ef)
        wrap = OxmlElement('wp:wrapSquare'); wrap.set('wrapText','bothSides'); anchor.append(wrap)
        if docPr       is not None: anchor.append(docPr)
        if cNvGrfFrmPr is not None: anchor.append(cNvGrfFrmPr)
        if graphic     is not None: anchor.append(graphic)

        drawing.remove(inline)
        drawing.append(anchor)

    # ── 圖片段落 ────────────────────────────────────────────

    def _download_image(self, src):
        """下載圖片並清除 EXIF，回傳 BytesIO 或 None"""
        try:
            req = urllib.request.Request(src, headers={'User-Agent': 'Mozilla/5.0'})
            img_bytes = urllib.request.urlopen(req, timeout=15).read()
            pil_img = PILImage.open(io.BytesIO(img_bytes))
            clean = io.BytesIO()
            fmt = (pil_img.format or 'JPEG').upper()
            pil_img.save(clean, format='JPEG' if fmt in ('JPEG','JPG') else 'PNG', quality=95)
            clean.seek(0)
            return clean
        except Exception as e:
            print(f'Warning: cannot download {src}: {e}')
            return None

    def _download_and_crop_square(self, src):
        """下載圖片，中心裁切為正方形，回傳 BytesIO 或 None"""
        try:
            req = urllib.request.Request(src, headers={'User-Agent': 'Mozilla/5.0'})
            img_bytes = urllib.request.urlopen(req, timeout=15).read()
            pil_img = PILImage.open(io.BytesIO(img_bytes))
            if pil_img.mode not in ('RGB', 'L'):
                pil_img = pil_img.convert('RGB')
            w, h = pil_img.size
            if w != h:
                size = min(w, h)
                left = (w - size) // 2
                top  = (h - size) // 2
                pil_img = pil_img.crop((left, top, left + size, top + size))
            clean = io.BytesIO()
            pil_img.save(clean, format='JPEG', quality=95)
            clean.seek(0)
            return clean
        except Exception as e:
            print(f'⚠️ 作者圖片載入失敗: {src}: {e}', file=sys.stderr)
            return None

    def _add_figure(self, html):
        src_m = re.search(r'src=["\']([^"\']+)["\']', html)
        alt_m = re.search(r'alt=["\']([^"\']*)["\']', html)
        cap_m = re.search(r'<figcaption>(.*?)</figcaption>', html, re.DOTALL)

        src     = src_m.group(1).strip() if src_m else ''
        alt     = alt_m.group(1).strip() if alt_m else ''
        caption_lines = []
        if cap_m:
            raw = cap_m.group(1)
            for part in re.split(r'<br\s*/?>', raw):
                part = re.sub(r'<[^>]+>', '', part).strip()
                if part:
                    caption_lines.append(part)

        float_dir = None
        if 'img-right' in html: float_dir = 'right'
        elif 'img-left' in html: float_dir = 'left'

        width_m  = re.search(r'px-(\d+)', html)
        px       = int(width_m.group(1)) if width_m else 250
        width_cm = min(px * 0.02, 6.0) if float_dir else min(px * 0.02, 14.2)

        # 偵測人物照片雙框線樣式（CSS border + outline）
        img_style_m = re.search(r'<img\b[^>]*style=["\']([^"\']*)["\']', html, re.IGNORECASE)
        has_portrait_border = bool(
            img_style_m and
            'outline' in img_style_m.group(1).lower() and
            'border'  in img_style_m.group(1).lower()
        )

        # 人物照片：正方形裁切 + 固定 7cm；一般圖片正常下載
        if has_portrait_border:
            img_stream = self._download_and_crop_square(src) if src else None
            width_cm = 7.0
        else:
            img_stream = self._download_image(src) if src else None

        if not img_stream:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'[圖片：{alt or src or "（未知）"}]')
            self._apply_font(run, 'Times New Roman', 'NSimSun', size=10, color=(0x80,0x80,0x80))
            return

        if float_dir:
            # 浮動圖（左/右）→ 用浮動表格呈現
            self._add_figure_textbox(img_stream, width_cm, caption_lines, float_dir,
                                     portrait_border=has_portrait_border)
        else:
            # 置中行內圖（img-bottom）→ 前後各空一行
            self._add_blank_line()
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            run = p.add_run()
            run.add_picture(img_stream, width=Cm(width_cm))
            if caption_lines:
                pc = self.doc.add_paragraph()
                pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pc.paragraph_format.space_before = Pt(0)
                pc.paragraph_format.space_after  = Pt(0)
                self._add_caption_runs(pc, caption_lines)
            self._add_blank_line()   # 圖說之後空一行

    def _add_figure_textbox(self, img_stream, width_cm, caption_lines, float_dir,
                            portrait_border=False):
        """浮動表格：圖片＋圖說組成矩形群組，整體文繞圖
        portrait_border=True 時加人物照片單層3px粗框（DrawingML，與行內圖一致）"""
        margin_dxa = 0
        table_width_dxa = int(width_cm * 567)

        tbl = self.doc.add_table(rows=1, cols=1)

        # ── tblPr ──
        tblPr = tbl._element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl._element.insert(0, tblPr)

        # 浮動定位
        tblpPr = OxmlElement('w:tblpPr')
        tblpPr.set(qn('w:leftFromText'),  '170')
        tblpPr.set(qn('w:rightFromText'), '170')
        tblpPr.set(qn('w:topFromText'),    '0')
        tblpPr.set(qn('w:bottomFromText'), '0')
        tblpPr.set(qn('w:vertAnchor'),  'text')
        tblpPr.set(qn('w:horzAnchor'), 'margin')
        tblpPr.set(qn('w:tblpXSpec'), float_dir)   # 'right' or 'left'
        tblPr.insert(0, tblpPr)

        # 表格寬度（含 margin）
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:w'),    str(table_width_dxa))
        tblW.set(qn('w:type'), 'dxa')

        # 表格框線（全部清除；portrait 的外框設在圖片 DrawingML）
        tblBorders = OxmlElement('w:tblBorders')
        for side in ('top','left','bottom','right','insideH','insideV'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'none')
            b.set(qn('w:sz'), '0')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
            tblBorders.append(b)
        tblPr.append(tblBorders)

        # 表格層級：強制 cellSpacing=0，防止 Word 預設留白影響框線位置
        tblCellSpacing = OxmlElement('w:tblCellSpacing')
        tblCellSpacing.set(qn('w:w'), '0')
        tblCellSpacing.set(qn('w:type'), 'dxa')
        tblPr.append(tblCellSpacing)

        tblCellMar = OxmlElement('w:tblCellMar')
        for side in ('top', 'left', 'bottom', 'right'):
            m = OxmlElement(f'w:{side}')
            m.set(qn('w:w'), '0')
            m.set(qn('w:type'), 'dxa')
            tblCellMar.append(m)
        tblPr.append(tblCellMar)

        # ── 儲存格 ──
        cell = tbl.rows[0].cells[0]
        tcPr = cell._tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            cell._tc.insert(0, tcPr)

        tcW = tcPr.find(qn('w:tcW'))
        if tcW is None:
            tcW = OxmlElement('w:tcW')
            tcPr.append(tcW)
        tcW.set(qn('w:w'),    str(table_width_dxa))
        tcW.set(qn('w:type'), 'dxa')

        tcBorders = OxmlElement('w:tcBorders')
        for side in ('top','left','bottom','right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'none')
            b.set(qn('w:sz'), '0')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
            tcBorders.append(b)
        tcPr.append(tcBorders)

        # 儲存格 margin：全部 0，圖片緊貼邊界
        tcMar = OxmlElement('w:tcMar')
        for side in ('top', 'left', 'bottom', 'right'):
            m = OxmlElement(f'w:{side}')
            m.set(qn('w:w'),    '0')
            m.set(qn('w:type'), 'dxa')
            tcMar.append(m)
        tcPr.append(tcMar)

        # 圖片（第 1 段）— 消除段前/後距，讓圖片貼緊儲存格邊界
        img_para = cell.paragraphs[0]
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.paragraph_format.space_before = Pt(0)
        img_para.paragraph_format.space_after  = Pt(0)
        img_run = img_para.add_run()
        if portrait_border:
            # 人物照片：寬高都是 7cm（已正方形裁切），套雙框線
            img_run.add_picture(img_stream, width=Cm(width_cm), height=Cm(width_cm))
            self._add_portrait_double_border(img_run)
        else:
            img_run.add_picture(img_stream, width=Cm(width_cm))

        # 圖說（第 2 段，支援 <br> 換行）
        if caption_lines:
            cap_para = cell.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_para.paragraph_format.space_before = Pt(0)
            cap_para.paragraph_format.space_after  = Pt(0)
            self._add_caption_runs(cap_para, caption_lines)

    def _add_caption_runs(self, paragraph, lines):
        """將圖說多行（原 <br> 分割）加入段落，行間用 w:br 換行"""
        for i, line in enumerate(lines):
            if i > 0:
                br_run = paragraph.add_run()
                br = OxmlElement('w:br')
                br_run._element.append(br)
            run = paragraph.add_run(line)
            self._apply_font(run, 'Times New Roman', 'PMingLiU', size=10, color=(0x59,0x59,0x59))

    # ── 結構元件 ─────────────────────────────────────────────

    def add_category_tag(self, category, color_hex):
        """欄目標籤：靠右，14pt，非粗體，無段前段後間距"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(f'【{category}】')
        self._apply_font(run, 'Times New Roman', '文鼎中行書', size=14, bold=False,
                         color=self._hex_rgb(color_hex))

    def add_title(self, title):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(title)
        self._apply_font(run, 'Times New Roman', '華康中黑體', size=24, bold=True)

    def add_subtitle(self, subtitle):
        if not subtitle:
            return
        if not subtitle.startswith('──'):
            subtitle = f'──{subtitle}'
        subtitle = f'\u3000\u3000{subtitle}'   # 前置兩個全形空格
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(subtitle)
        self._apply_font(run, 'Times New Roman', '華康中黑體', size=16, bold=True)

    def add_decoration_line(self):
        """裝飾線：單段落，上框 3px（粗）＋下框 1px（細），行高 1pt 讓兩線間距約 1px"""
        p = self.doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '0')
        sp.set(qn('w:line'), '20');  sp.set(qn('w:lineRule'), 'exact')  # 1pt 行高
        pPr.append(sp)
        pBdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '18')   # ~3px
        top.set(qn('w:space'), '0');    top.set(qn('w:color'), '000000')
        pBdr.append(top)
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')    # ~1px
        bot.set(qn('w:space'), '0');    bot.set(qn('w:color'), '000000')
        pBdr.append(bot)
        pPr.append(pBdr)

    def add_author(self, author, author_title=None, remark=None):
        def _add_field(text, ascii_font, east_font, size, allow_footnote=False):
            if not text:
                return
            parts = re.split(r'<br\s*/?>', text, flags=re.IGNORECASE)
            for part in parts:
                part = part.strip()
                if not part:
                    continue
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.space_before    = Pt(0)
                p.paragraph_format.space_after     = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                p.paragraph_format.line_spacing      = 1.5
                if allow_footnote:
                    # 支援 [^N] → Word 腳注引用；其餘 HTML 剝除後以備註字體輸出
                    fn_pat = re.compile(r'\[\^(\d+)\]')
                    segs = fn_pat.split(part)
                    for i, chunk in enumerate(segs):
                        if i % 2 == 1:
                            # 奇數位是腳注編號
                            self._add_footnote_ref(p, chunk)
                        else:
                            plain = re.sub(r'<[^>]+>', '', chunk).strip()
                            if plain:
                                run = p.add_run(plain)
                                self._apply_font(run, ascii_font, east_font, size=size)
                else:
                    plain = re.sub(r'<[^>]+>', '', part).strip()
                    if plain:
                        run = p.add_run(plain)
                        self._apply_font(run, ascii_font, east_font, size=size)

        _add_field(author,       'Brush Script MT', '文鼎中行書', 12)
        _add_field(author_title, 'Brush Script MT', '文鼎中行書', 12)
        _add_field(remark,       'Brush Script MT', '文鼎中行書', 12, allow_footnote=True)

    def add_keywords(self, keywords):
        if not keywords:
            return
        kw = re.sub(r'^[🌿\s]*(關鍵字[：:])?\s*', '', keywords).strip()
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        for chunk, is_emoji in self._split_emoji(f'🌿關鍵字：{kw}'):
            run = p.add_run(chunk)
            if is_emoji:
                self._apply_font(run, 'Segoe UI Symbol', 'Segoe UI Symbol', size=12)
            else:
                self._apply_font(run, 'Times New Roman', 'NSimSun', size=12, bold=True)

    # ── 內文解析 ─────────────────────────────────────────────

    def _split_into_segments(self, content):
        """將 HTML 內容切分為 (type, html) segments，正確處理嵌套標籤。"""
        BLOCK_TAGS = ('figure', 'blockquote', 'table', 'div', 'p', 'h1', 'h2', 'h3')
        start_re = re.compile(
            r'<(' + '|'.join(BLOCK_TAGS) + r')(\b[^>]*)?>',
            re.IGNORECASE
        )
        segments = []
        pos = 0
        while pos < len(content):
            m = start_re.search(content, pos)
            if not m:
                tail = content[pos:].strip()
                if tail:
                    segments.append(('text', tail))
                break
            before = content[pos:m.start()].strip()
            if before:
                segments.append(('text', before))
            tag = m.group(1).lower()
            end_abs = self._find_close_tag(content, m.end(), tag)
            segments.append((tag, content[m.start():end_abs]))
            pos = end_abs
        return segments

    def _find_close_tag(self, content, from_pos, tag):
        """找到對應的關閉標籤位置（處理嵌套），回傳關閉標籤結束後的位置。"""
        open_re  = re.compile(r'<'  + re.escape(tag) + r'\b', re.IGNORECASE)
        close_re = re.compile(r'</' + re.escape(tag) + r'\s*>', re.IGNORECASE)
        depth = 1
        pos   = from_pos
        while pos < len(content) and depth > 0:
            om = open_re.match(content, pos)
            cm = close_re.match(content, pos)
            if om:
                depth += 1
                gt = content.find('>', pos)
                pos = (gt + 1) if gt != -1 else pos + 1
            elif cm:
                depth -= 1
                pos = cm.end()
            else:
                pos += 1
        return pos

    def add_content(self, content):
        if not content:
            return
        content = re.sub(r'(\s*<br\s*/?>)+\s*$', '', content.strip())
        for seg_type, seg_html in self._split_into_segments(content):
            if seg_type == 'figure':
                self._add_figure(seg_html)
            elif seg_type == 'blockquote':
                inner = seg_html
                # 先抓出 <div class="rel"> 作為來源行（置右）
                rel_m = re.search(
                    r'<div[^>]*class=["\'][^"\']*\brel\b[^"\']*["\'][^>]*>(.*?)</div>',
                    inner, re.DOTALL | re.IGNORECASE)
                rel_text = ''
                if rel_m:
                    rel_text = re.sub(r'<[^>]+>', '', rel_m.group(1)).strip()
                    inner = inner[:rel_m.start()] + inner[rel_m.end():]
                # 換行語義：<br> / </p> → \n
                inner = re.sub(r'<br\s*/?>', '\n', inner, flags=re.IGNORECASE)
                inner = re.sub(r'</p\s*>', '\n', inner, flags=re.IGNORECASE)
                # 只剝除塊級結構標籤；保留 inline 格式標籤供 _add_inline 解析
                inner = re.sub(
                    r'</?(?:blockquote|p|div|section|article|header|footer|ul|ol|li)\b[^>]*>',
                    '', inner, flags=re.IGNORECASE | re.DOTALL)
                inner = inner.strip()
                if inner or rel_text:
                    self._add_blockquote(inner, rel_text=rel_text)
            elif seg_type == 'div':
                self._dispatch_div(seg_html)
            elif seg_type == 'table':
                self._add_table(seg_html)
            elif seg_type == 'p':
                self._add_paragraph_element(seg_html)
            elif seg_type in ('h1', 'h2', 'h3'):
                inner = re.sub(r'^<h[1-3][^>]*>', '', seg_html, count=1, flags=re.IGNORECASE)
                inner = re.sub(r'</h[1-3]>\s*$', '', inner, flags=re.IGNORECASE).strip()
                if inner:
                    self._add_section_title(inner)
            else:  # text
                normalized = re.sub(r'<br\s*/?>', '\n', seg_html)
                for line in normalized.split('\n'):
                    self._process_line(line.strip())

    def _dispatch_div(self, html):
        """依 class 屬性將 <div> 派發至對應處理器。"""
        outer_class = ''
        m = re.match(r'<div\b[^>]*class=["\']([^"\']*)["\']', html, re.IGNORECASE)
        if m:
            outer_class = m.group(1)

        if 'custom-divider' in outer_class:
            self._add_custom_divider()
        elif 'book-quote' in outer_class:
            self._add_book_quote(html)
        elif 'book-box' in outer_class:
            self._add_book_box(html)
        elif 'reference-box' in outer_class:
            self._add_reference_box(html)
        elif 'theme-image' in outer_class:
            self._add_theme_image(html)
        elif 'author-profile' in outer_class:
            self._add_author_profile(html)
        elif 'footnotes' in outer_class:
            pass  # 腳注由 _finalize_footnotes 處理
        else:
            # 一般 div：strip 所有 HTML，取出文字後輸出段落
            style_m = re.search(r'style=["\']([^"\']+)["\']', html, re.IGNORECASE)
            inner = re.sub(r'<br\s*/?>', '\n', html)
            inner = re.sub(r'<[^>]+>', '', inner, flags=re.DOTALL).strip()
            if not inner:
                return
            if style_m and 'text-align' in style_m.group(1).lower():
                if 'right' in style_m.group(1).lower():
                    self._add_right_aligned(inner)
                    return
            for line in inner.split('\n'):
                line = line.strip()
                if line:
                    p = self.doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Pt(24)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after  = Pt(0)
                    self._add_inline(p, line)

    def _add_paragraph_element(self, html):
        """處理 <p> 元素，支援 no-indent、toc-line、text-align:right 與內部 <br>。
        <p class="no-indent"> 前自動插入一空行。
        <p class="toc-line"> 凸排：左縮排 24pt，首行 -24pt。"""
        no_indent = bool(re.search(r'class=["\'][^"\']*no-indent[^"\']*["\']',
                                   html, re.IGNORECASE))
        toc_line  = bool(re.search(r'class=["\'][^"\']*toc-line[^"\']*["\']',
                                   html, re.IGNORECASE))
        # 偵測 text-align: right → 交由 _add_right_aligned 統一處理
        tag_m = re.match(r'<p\b([^>]*)>', html, re.IGNORECASE)
        if tag_m:
            style_m = re.search(r'style=["\']([^"\']*)["\']', tag_m.group(1), re.IGNORECASE)
            if style_m and 'text-align' in style_m.group(1) and 'right' in style_m.group(1):
                inner = re.sub(r'^<p[^>]*>', '', html, count=1, flags=re.IGNORECASE)
                inner = re.sub(r'</p>\s*$', '', inner, flags=re.IGNORECASE).strip()
                inner = re.sub(r'<[^>]+>', '', inner).strip()
                if inner:
                    self._add_right_aligned(inner)
                return
        inner = re.sub(r'^<p[^>]*>', '', html, count=1, flags=re.IGNORECASE)
        inner = re.sub(r'</p>\s*$',  '', inner, flags=re.IGNORECASE).strip()
        if not inner:
            return
        # <p class="no-indent"> 前空一行
        # —— 但前面已是空行或小標題時跳過，避免重複
        if no_indent and not self._should_skip_blank_before():
            self._add_blank_line()
        parts = re.split(r'<br\s*/?>', inner, flags=re.IGNORECASE)
        toc_first = True  # 第一個 part 才凸排，後續 parts 只左縮排
        for part in parts:
            part = part.strip()
            if not part:
                continue
            bullet_m = re.match(r'&#9679;(?:&nbsp;|\u00a0|\s)*', part, re.IGNORECASE)
            if bullet_m:
                self._add_bullet_line(part[bullet_m.end():])
                continue
            p = self.doc.add_paragraph()
            if toc_line:
                if toc_first:
                    # 第一行凸排：左縮排 18pt，首行 -18pt（從 0 開始）
                    p.paragraph_format.left_indent       = Pt(18)
                    p.paragraph_format.first_line_indent = Pt(-18)
                    toc_first = False
                else:
                    # 後續行（<br> 分段）從 18pt 對齊
                    p.paragraph_format.left_indent       = Pt(18)
                    p.paragraph_format.first_line_indent = Pt(0)
            elif no_indent:
                p.paragraph_format.first_line_indent = Pt(0)
            else:
                p.paragraph_format.first_line_indent = Pt(24)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            if toc_line:
                # 目次行格式：序號+分類 12pt，文章主標題 13pt，副標題以後 12pt
                if '──' in part:
                    # 副標題行（<br> 之後），全部 12pt
                    self._add_inline(p, part, size=12)
                else:
                    # 主標題行：找最後一個 HTML 結束 > 的位置
                    # 之前 = 序號 + 分類標籤（12pt），之後 = 純標題文字（13pt）
                    last_gt = max((m.end() for m in re.finditer(r'>', part)), default=0)
                    if last_gt > 0:
                        self._add_inline(p, part[:last_gt], size=12)
                        self._add_inline(p, part[last_gt:], size=13)
                    else:
                        self._add_inline(p, part, size=13)
            else:
                self._add_inline(p, part)

    def _add_custom_divider(self):
        """灰色水平分隔線，對應 CSS .custom-divider（3px gray）。
        前後各一空行；前面已是空行或小標題時略過前置空行。"""
        if not self._should_skip_blank_before():
            self._add_blank_line()
        p = self.doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        sp_el = OxmlElement('w:spacing')
        sp_el.set(qn('w:before'), '0')
        sp_el.set(qn('w:after'),  '0')
        sp_el.set(qn('w:line'),   '40')
        sp_el.set(qn('w:lineRule'), 'exact')
        pPr.append(sp_el)
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '18')      # 18/8 = 2.25pt ≈ 3px
        bot.set(qn('w:space'), '0')
        bot.set(qn('w:color'), 'AAAAAA')  # 灰色
        pBdr.append(bot)
        pPr.append(pBdr)
        self._add_blank_line()

    def _add_book_quote(self, html):
        """書本引言（.book-quote）：咖啡色左豎線，楷書體，右對齊出處行。"""
        rel_m = re.search(r'<div[^>]*book-quote-rel[^>]*>(.*?)</div>',
                          html, re.DOTALL | re.IGNORECASE)
        rel_text = ''
        if rel_m:
            rel_text = re.sub(r'<[^>]+>', '', rel_m.group(1)).strip()
            html = html[:rel_m.start()] + html[rel_m.end():]

        inner = re.sub(r'^<div[^>]*>', '', html, count=1, flags=re.IGNORECASE).strip()
        inner = re.sub(r'</div>\s*$',  '', inner, flags=re.IGNORECASE).strip()
        inner = re.sub(r'<br\s*/?>', '\n', inner, flags=re.IGNORECASE)
        lines = [re.sub(r'<[^>]+>', '', l).strip() for l in inner.split('\n')]
        lines = [l for l in lines if l]

        self._add_blank_line()  # 前置空行

        def _bq_para(space_after=0):
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent       = Pt(24)
            p.paragraph_format.right_indent      = Pt(24)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(space_after)
            pPr = p._element.get_or_add_pPr()
            # 行距 1.5 倍
            sp = OxmlElement('w:spacing')
            sp.set(qn('w:before'), '0')
            sp.set(qn('w:after'),  str(int(space_after * 20)))
            sp.set(qn('w:line'),   '360')   # 240 × 1.5
            sp.set(qn('w:lineRule'), 'auto')
            pPr.append(sp)
            # 左邊線：藍色 3px（18/8 = 2.25pt ≈ 3px）
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'),   'single')
            left.set(qn('w:sz'),    '18')
            left.set(qn('w:space'), '4')
            left.set(qn('w:color'), '1E6FD9')   # 藍色
            pBdr.append(left)
            pPr.append(pBdr)
            return p

        for line in lines:
            p = _bq_para()
            self._add_inline(p, line,
                             east_font='文鼎粗鋼筆行楷',
                             ascii_font='Brush Script MT')

        if rel_text:
            p_rel = _bq_para(space_after=6)
            p_rel.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self._add_inline(p_rel, rel_text,
                             east_font='文鼎粗鋼筆行楷',
                             ascii_font='Brush Script MT')

        self._add_blank_line()  # 後置空行

    def _add_book_box(self, html):
        """書籍簡介框（.book-box）：2 欄無框線表格，文字（2/3）+ 封面圖（1/3）。"""
        info_m = re.search(r'<div[^>]*book-info[^>]*>(.*?)</div>',
                           html, re.DOTALL | re.IGNORECASE)
        img_m  = re.search(r'<img[^>]+src=["\']([^"\']+)["\']', html, re.IGNORECASE)
        if not info_m:
            return

        info_html = info_m.group(1)
        img_src   = img_m.group(1) if img_m else None

        self._add_blank_line()

        # ── 建立 2 欄無框線表格 ──────────────────────────────────────
        PAGE_W = 8051           # 14.2cm（頁面可用寬度）in twips
        TEXT_W = PAGE_W * 2 // 3
        IMG_W  = PAGE_W - TEXT_W

        tbl    = self.doc.add_table(rows=1, cols=2)
        tbl_el = tbl._tbl
        tblPr  = tbl_el.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl_el.insert(0, tblPr)

        # 移除預設表格樣式
        for old in tblPr.findall(qn('w:tblStyle')):
            tblPr.remove(old)

        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'),    str(PAGE_W))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)

        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)

        tblBorders = OxmlElement('w:tblBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'none')
            tblBorders.append(b)
        tblPr.append(tblBorders)

        row   = tbl.rows[0]
        l_cel = row.cells[0]
        r_cel = row.cells[1]

        for cell, w in ((l_cel, TEXT_W), (r_cel, IMG_W)):
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(qn('w:tcW')):
                tcPr.remove(old)
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'),    str(w))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            tcBdr = OxmlElement('w:tcBdr')
            for side in ('top', 'left', 'bottom', 'right'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'none')
                tcBdr.append(b)
            tcPr.append(tcBdr)
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)

        # ── 文字欄：綠色左豎線 + _add_inline（保留粗體/斜體） ────────
        def _bb_para(cell, first=False):
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            p.paragraph_format.left_indent       = Pt(12)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(3)
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            lb = OxmlElement('w:left')
            lb.set(qn('w:val'),   'single')
            lb.set(qn('w:sz'),    '20')
            lb.set(qn('w:space'), '4')
            lb.set(qn('w:color'), '378B13')
            pBdr.append(lb)
            pPr.append(pBdr)
            return p

        info_lines = re.sub(r'<br\s*/?>', '\n', info_html, flags=re.IGNORECASE).split('\n')
        first = True
        for raw in info_lines:
            raw = raw.strip()
            if not raw:
                continue
            p = _bb_para(l_cel, first=first)
            first = False
            self._add_inline(p, raw, east_font='標楷體')

        # ── 圖片欄：置中，寬度填滿欄位 ──────────────────────────────
        if img_src:
            img_stream = self._download_image(img_src)
            if img_stream:
                img_cm = round(IMG_W / 567 - 0.4, 2)
                img_p  = r_cel.paragraphs[0]
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.paragraph_format.space_before      = Pt(0)
                img_p.paragraph_format.space_after       = Pt(0)
                img_p.paragraph_format.first_line_indent = Pt(0)
                img_p.add_run().add_picture(img_stream, width=Cm(img_cm))

        self._add_blank_line()

    def _add_reference_box(self, html):
        """參考資料框（.reference-box）：綠色左豎線，條列來源。"""
        strong_m = re.search(r'<strong>(.*?)</strong>', html, re.DOTALL | re.IGNORECASE)
        title_text = re.sub(r'<[^>]+>', '', strong_m.group(1)).strip() if strong_m else '參考資料'

        inner = html
        if strong_m:
            inner = inner[:strong_m.start()] + inner[strong_m.end():]
        inner = re.sub(r'^<div[^>]*>', '', inner, count=1, flags=re.IGNORECASE).strip()
        inner = re.sub(r'</div>\s*$',  '', inner, flags=re.IGNORECASE).strip()
        inner = re.sub(r'<br\s*/?>', '\n', inner)
        body_text = re.sub(r'<[^>]+>', '', inner).strip()

        self._add_blank_line()

        def _ref_para():
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent       = Pt(24)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'),   'single')
            left.set(qn('w:sz'),    '36')
            left.set(qn('w:space'), '4')
            left.set(qn('w:color'), '378B13')
            pBdr.append(left)
            pPr.append(pBdr)
            return p

        p_title = _ref_para()
        run = p_title.add_run(title_text)
        self._apply_font(run, 'Times New Roman', 'NSimSun', size=13, bold=True)

        for line in body_text.split('\n'):
            line = line.strip()
            if line:
                p = _ref_para()
                run = p.add_run('• ' + line)
                self._apply_font(run, 'Times New Roman', 'NSimSun', size=11)

        self._add_blank_line()

    def _add_theme_image(self, html):
        """主題圖片（.theme-image）：置中大圖，無圖說。"""
        src_m = re.search(r'src=["\']([^"\']+)["\']', html)
        if not src_m:
            return
        src = src_m.group(1).strip()
        img_stream = self._download_image(src)
        if not img_stream:
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        p.add_run().add_picture(img_stream, width=Cm(12.0))

    def _add_author_profile(self, html):
        """作者簡介（.author-profile）：左圖右文表格排版，符合原版 PDF 格式。
        左欄：5×5cm 正方形照片（含 2px 邊框）；右欄：粗體姓名 + 簡介段落。"""
        # 提取圖片網址
        img_m = re.search(r'<img\b[^>]*\bsrc=["\']([^"\']+)["\']', html, re.IGNORECASE)
        img_src = img_m.group(1).strip() if img_m else None

        # 提取作者姓名（h3）
        name_m = re.search(r'<h3[^>]*>(.*?)</h3>', html, re.IGNORECASE | re.DOTALL)
        name = re.sub(r'<[^>]+>', '', name_m.group(1)).strip() if name_m else ''

        # 提取作者簡介（所有 p 段落）
        bio_parts = re.findall(r'<p[^>]*>(.*?)</p>', html, re.IGNORECASE | re.DOTALL)
        bio_lines = [re.sub(r'<[^>]+>', '', b).strip() for b in bio_parts]
        bio = '\n'.join(l for l in bio_lines if l)

        # 建立 2 欄表格（左欄圖片 5.5cm，右欄文字 8.7cm）
        tbl = self.doc.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'

        # 表格寬度 = 頁面可用寬度 18.2 - 2 - 2 = 14.2cm = 8051 twips，靠左對齊
        tbl_el = tbl._tbl
        tblPr = tbl_el.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl_el.insert(0, tblPr)
        for old in tblPr.findall(qn('w:tblW')):
            tblPr.remove(old)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'),    '8051')   # 14.2cm in twips
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)

        row = tbl.rows[0]
        left_cell  = row.cells[0]
        right_cell = row.cells[1]

        # 設定欄寬（1cm = 567 twips）並在每個 cell 上設邊框顏色（覆蓋 Table Grid 樣式）
        for cell, w_dxa in ((left_cell, int(5.5 * 567)), (right_cell, int(8.7 * 567))):
            tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn('w:tcW')):
                tcPr.remove(old)
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'),    str(w_dxa))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            # 每格邊框：黑色淺15%（#262626），cell 層級優先於 style
            for old in tcPr.findall(qn('w:tcBorders')):
                tcPr.remove(old)
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                bd = OxmlElement(f'w:{side}')
                bd.set(qn('w:val'),   'single')
                bd.set(qn('w:sz'),    '4')
                bd.set(qn('w:space'), '0')
                bd.set(qn('w:color'), '262626')   # 黑色淺15%
                tcBorders.append(bd)
            tcPr.append(tcBorders)

        # 左欄：垂直置中 + 上下 5px 內距（5px ≈ 75 twips at 96dpi）
        left_tcPr = left_cell._tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        left_tcPr.append(vAlign)
        tcMar = OxmlElement('w:tcMar')
        for side in ('top', 'bottom'):
            mar = OxmlElement(f'w:{side}')
            mar.set(qn('w:w'),    '266')   # 226(15px) + 40(2pt border) = 266 twips
            mar.set(qn('w:type'), 'dxa')
            tcMar.append(mar)
        left_tcPr.append(tcMar)

        # 左欄：照片（5×5cm，正方形裁切，2pt 邊框）
        lp = left_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after  = Pt(0)
        if img_src:
            img_stream = self._download_and_crop_square(img_src)
            if img_stream:
                img_run = lp.add_run()
                img_run.add_picture(img_stream, width=Cm(5.0), height=Cm(5.0))
                self._add_picture_border(img_run, border_pt=2.0)
            else:
                lp.add_run(f'[圖片：{img_src}]')

        # 右欄：作者姓名（粗體）+ 簡介
        rp_name = right_cell.paragraphs[0]
        rp_name.paragraph_format.space_before = Pt(6)
        rp_name.paragraph_format.space_after  = Pt(4)
        if name:
            name_run = rp_name.add_run(name)
            self._apply_font(name_run, 'Times New Roman', 'NSimSun', size=12, bold=True)

        if bio:
            for bio_line in bio.split('\n'):
                bio_line = bio_line.strip()
                if not bio_line:
                    continue
                rp_bio = right_cell.add_paragraph()
                rp_bio.paragraph_format.space_before = Pt(0)
                rp_bio.paragraph_format.space_after  = Pt(0)
                bio_run = rp_bio.add_run(bio_line)
                self._apply_font(bio_run, 'Times New Roman', 'NSimSun', size=12)

        # 簡介後空一行
        rp_blank = right_cell.add_paragraph()
        rp_blank.paragraph_format.space_before = Pt(0)
        rp_blank.paragraph_format.space_after  = Pt(0)

    def _add_right_aligned(self, text):
        """置右段落（對應 style="text-align: right"）。"""
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            # 「作者依照文章先後順序排列」：新細明體 10pt，黑色淺15%（#262626），置右
            clean_line = re.sub(r'<[^>]+>', '', line).strip()
            if '作者依照文章先後順序排列' in clean_line:
                run = p.add_run(clean_line)
                self._apply_font(run, 'PMingLiU', 'PMingLiU', size=10,
                                 color=(38, 38, 38))
            else:
                self._add_inline(p, line)

    def _add_table(self, html):
        """解析並插入 HTML 表格（<table class="data-table">）。"""
        rows_html = re.findall(r'<tr\b[^>]*>(.*?)</tr>', html,
                               re.DOTALL | re.IGNORECASE)
        if not rows_html:
            return
        first_cells = re.findall(r'<(?:td|th)\b[^>]*>.*?</(?:td|th)>',
                                 rows_html[0], re.DOTALL | re.IGNORECASE)
        num_cols = max(len(first_cells), 1)
        tbl = self.doc.add_table(rows=0, cols=num_cols)
        tbl.style = 'Table Grid'
        for row_html in rows_html:
            cells = re.findall(r'<(td|th)\b[^>]*>(.*?)</(?:td|th)>',
                               row_html, re.DOTALL | re.IGNORECASE)
            row = tbl.add_row()
            for i, (ctag, ccontent) in enumerate(cells[:num_cols]):
                p = row.cells[i].paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                text = re.sub(r'<[^>]+>', '', ccontent).strip()
                if text:
                    run = p.add_run(text)
                    self._apply_font(run, 'Times New Roman', 'NSimSun',
                                     size=11, bold=(ctag.lower() == 'th'))
        self._add_blank_line()

    def _add_portrait_double_border(self, run):
        """人物照片雙框線：1pt 細線 + 1pt 空格 + 3pt 粗線（DrawingML thinThick compound）"""
        from lxml import etree as lxml_etree
        PIC_NS = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
        A_NS   = 'http://schemas.openxmlformats.org/drawingml/2006/main'

        drawing = run._r.find(qn('w:drawing'))
        if drawing is None:
            return
        pic_el = drawing.find('.//' + f'{{{PIC_NS}}}pic')
        if pic_el is None:
            return
        spPr_el = pic_el.find(f'{{{PIC_NS}}}spPr')
        if spPr_el is None:
            spPr_el = lxml_etree.SubElement(pic_el, f'{{{PIC_NS}}}spPr')

        for ln_el in list(spPr_el.findall(f'{{{A_NS}}}ln')):
            spPr_el.remove(ln_el)
        for geom_el in list(spPr_el.findall(f'{{{A_NS}}}prstGeom')):
            spPr_el.remove(geom_el)

        # 明確宣告矩形形狀，框線才不會被截斷
        geom_el = lxml_etree.SubElement(
            spPr_el, f'{{{A_NS}}}prstGeom', attrib={'prst': 'rect'})
        lxml_etree.SubElement(geom_el, f'{{{A_NS}}}avLst')

        # thinThick：細線（內）+ 自動空格 + 粗線（外）
        # 總寬 ≈ 5pt → 細線≈1pt、空格≈1pt、粗線≈3pt
        total_emu = int(5 * 12700)
        ln_el = lxml_etree.SubElement(
            spPr_el, f'{{{A_NS}}}ln',
            attrib={'w': str(total_emu), 'cap': 'flat', 'cmpd': 'thinThick', 'algn': 'ctr'})
        sf_el = lxml_etree.SubElement(ln_el, f'{{{A_NS}}}solidFill')
        lxml_etree.SubElement(sf_el, f'{{{A_NS}}}srgbClr', attrib={'val': '000000'})
        lxml_etree.SubElement(ln_el, f'{{{A_NS}}}prstDash', attrib={'val': 'solid'})

    def _add_picture_border(self, run, border_pt=0.75, color='000000'):
        """在已插入圖片的 run 上加 drawingML 內框線（模擬 CSS border）。"""
        PIC_NS = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
        A_NS   = 'http://schemas.openxmlformats.org/drawingml/2006/main'
        from lxml import etree as lxml_etree

        drawing = run._r.find(qn('w:drawing'))
        if drawing is None:
            return
        pic_el = drawing.find('.//' + f'{{{PIC_NS}}}pic')
        if pic_el is None:
            return
        spPr_el = pic_el.find(f'{{{PIC_NS}}}spPr')
        if spPr_el is None:
            spPr_el = lxml_etree.SubElement(pic_el, f'{{{PIC_NS}}}spPr')

        # 移除舊框線
        for ln_el in list(spPr_el.findall(f'{{{A_NS}}}ln')):
            spPr_el.remove(ln_el)

        width_emu = int(border_pt * 12700)  # pt → EMU
        ln_el = lxml_etree.SubElement(
            spPr_el, f'{{{A_NS}}}ln',
            attrib={'w': str(width_emu), 'cap': 'flat', 'cmpd': 'sng', 'algn': 'ctr'})
        sf_el = lxml_etree.SubElement(ln_el, f'{{{A_NS}}}solidFill')
        lxml_etree.SubElement(sf_el, f'{{{A_NS}}}srgbClr', attrib={'val': color})
        lxml_etree.SubElement(ln_el, f'{{{A_NS}}}prstDash', attrib={'val': 'solid'})

    def _should_skip_blank_before(self):
        """若前一個段落已是空行或小標題，回傳 True，避免重複插入空行。"""
        paras = self.doc.paragraphs
        if not paras:
            return True
        last = paras[-1]
        # 前一段是空行（空白或無文字）
        if not last.text.strip():
            return True
        # 前一段是小標題（粗體，字型 ≥ 13pt）
        for run in last.runs:
            if run.font.bold and run.font.size and run.font.size >= Pt(13):
                return True
        return False

    def _add_bullet_line(self, text):
        """項目列表行：小圓點 • + 左縮排兩個字（24pt）"""
        text = _html_mod.unescape(text)
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent       = Pt(24)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        bullet_run = p.add_run('• ')
        self._apply_font(bullet_run, 'Times New Roman', 'NSimSun', size=11)
        self._add_inline(p, text)

    def _process_line(self, line):
        """逐行處理正文（跳過空行、處理標題與一般段落）"""
        if not line:
            return  # 略過空行，不插入空段落

        # 偵測 &#9679;（●）開頭的項目列表 → 改為小圓點 + 縮排兩個字
        bullet_m = re.match(r'&#9679;(?:&nbsp;|\u00a0|\s)*', line, re.IGNORECASE)
        if bullet_m:
            self._add_bullet_line(line[bullet_m.end():])
            return

        heading_match = re.match(r'^(#{1,3})\s+(.+)', line)
        if heading_match:
            self._add_section_title(heading_match.group(2))
            return

        if line.startswith('>'):
            self._add_blockquote(line[1:].strip())
            return

        # 偵測 <p class="no-indent"> 包裝 → 無首行縮排 + 前空一行
        no_indent = False
        no_indent_m = re.match(
            r'<p\s[^>]*class=[^>]*no-indent[^>]*>(.*)</p>$',
            line, re.IGNORECASE | re.DOTALL)
        if no_indent_m:
            line = no_indent_m.group(1).strip()
            no_indent = True
            # no-indent 段落前空一行（前面已是空行或小標題則跳過）
            if not self._should_skip_blank_before():
                self._add_blank_line()

        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0) if no_indent else Pt(24)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        self._add_inline(p, line)

    def _add_section_title(self, inner_html):
        """段落小標題：14pt 粗體，前置空行，段後 9pt（0.5 行距），無首行縮排，無列表符號
        inner_html 可包含 inline 標記（如 <sup class="footnote-ref">）。
        """
        import unicodedata
        # 去掉內容前置的符號/標點字元（■ ● ▪ 等），但保留裝飾性符號 ☆ ◇ ★ ◆
        KEEP_SYMBOLS = set('☆◇★◆')
        # 只對開頭的純文字部分做符號剝除（標籤之前）
        leading_text_m = re.match(r'^([^<]*)', inner_html)
        leading_text = leading_text_m.group(1) if leading_text_m else ''
        strip_count = 0
        while strip_count < len(leading_text):
            ch = leading_text[strip_count]
            if ch in KEEP_SYMBOLS or unicodedata.category(ch) not in (
                    'So', 'Sm', 'Sk', 'Sc', 'Po', 'Ps', 'Pe', 'Pi', 'Pf', 'Pd', 'Pc', 'Zs'):
                break
            strip_count += 1
        inner_html = (leading_text[strip_count:] + inner_html[len(leading_text):]).strip()
        if not inner_html:
            return

        # 前置空行（前面已是空行或小標題時略過，避免重複）
        if not self._should_skip_blank_before():
            self._add_blank_line()

        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 明確覆蓋 Normal style 的 1.5× 行距，改為單行；前後各 0.75 行距
        pPr = p._element.get_or_add_pPr()
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:beforeLines'), '75')    # 0.75 行距
        sp.set(qn('w:afterLines'),  '75')    # 0.75 行距
        sp.set(qn('w:line'),        '240')   # 單行行距（240 = 1×）
        sp.set(qn('w:lineRule'),    'auto')
        pPr.append(sp)

        # 逐 token 處理 inline 標記；一般文字 14pt 粗體，腳注引用則插入 Word footnote
        H3_TOKEN = re.compile(
            r'(<sup\b[^>]*class="footnote-ref"[^>]*><a\b[^>]*>[^<]*</a></sup>'
            r'|\[\^\d+\]'
            r'|<[^>]+>)',
            re.IGNORECASE | re.DOTALL)
        for seg in H3_TOKEN.split(inner_html):
            if not seg:
                continue
            fn_sup_m = re.search(r'<a\b[^>]*>([^<]*)</a>', seg) if seg.startswith('<sup') else None
            fn_sup   = fn_sup_m if (fn_sup_m and fn_sup_m.group(1).strip().isdigit()) else None
            fn_m     = re.fullmatch(r'\[\^(\d+)\]', seg)
            html_tag = re.fullmatch(r'<[^>]+>', seg)

            if fn_sup:
                self._add_footnote_ref(p, fn_sup.group(1))
            elif fn_m:
                self._add_footnote_ref(p, fn_m.group(1))
            elif html_tag:
                pass  # 略過其他 HTML 標籤
            else:
                text = _html_mod.unescape(seg)
                if text:
                    run = p.add_run(text)
                    self._apply_font(run, 'Times New Roman', 'NSimSun', size=14, bold=True)

    def _add_blockquote(self, text, rel_text=''):
        """獨立引用：標楷體，左縮排 24pt，首行與末行各空一行，中間行無間距。
        rel_text 為來源行（如 ──但丁），置右輸出。"""
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        last_is_rel = bool(rel_text)
        for i, line in enumerate(lines):
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent       = Pt(24)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(12) if i == 0 else Pt(0)
            p.paragraph_format.space_after  = Pt(12) if (i == len(lines) - 1 and not last_is_rel) else Pt(0)
            self._add_inline(p, line, east_font='標楷體')
        if rel_text:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent       = Pt(24)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self._add_inline(p, rel_text, east_font='標楷體')

    def _split_special_chars(self, text):
        """拆分含 ‧／∕ 的文字為 [(chunk, is_special), ...]，特殊字元用新細明體"""
        SPECIAL = re.compile(r'([‧／∕]+)')
        parts = SPECIAL.split(text)
        result = []
        for i, part in enumerate(parts):
            if part:
                result.append((part, i % 2 == 1))
        return result if result else [(text, False)]

    def _split_emoji(self, text):
        """將文字拆分為 [(chunk, is_emoji), ...] 以便分別套用字型"""
        result = []
        current = ''
        current_emoji = False
        for c in text:
            cp = ord(c)
            is_e = (cp > 0xFFFF or (0x2600 <= cp <= 0x27BF) or
                    (0x1F000 <= cp <= 0x1FFFF) or cp in (0xFE0F, 0x200D))
            if is_e != current_emoji:
                if current:
                    result.append((current, current_emoji))
                current = c
                current_emoji = is_e
            else:
                current += c
        if current:
            result.append((current, current_emoji))
        return result

    def _add_inline(self, paragraph, text, east_font='NSimSun', ascii_font='Times New Roman', size=12):
        """
        解析 inline 標記：
          **text**           → 粗體
          *text*             → 標楷體（行內引用）
          <b>/<strong>       → 粗體
          <em>               → 標楷體
          <i>                → 斜體
          <br>               → 段落內換行（w:br）
          <span style="..."> → 小字體 / 置右等樣式
          [^N]               → Word 腳注引用
          其他 <tag>          → 略過（不輸出原始文字）
          emoji              → Segoe UI Symbol
        """
        TOKEN = re.compile(
            r'(\*\*[^*\n]+?\*\*'
            r'|\*[^*\n]+?\*'
            r'|<(?:b|strong)>[^<]*</(?:b|strong)>'
            r'|<em>[^<]*</em>'
            r'|<i>[^<]*</i>'
            r'|<br\s*/?>'
            r'|<span\b[^>]*>.*?</span>'
            r'|<sup\b[^>]*class="footnote-ref"[^>]*><a\b[^>]*>[^<]*</a></sup>'
            r'|<[^>]+>'
            r'|\[\^\d+\])',
            re.IGNORECASE | re.DOTALL)
        for seg in TOKEN.split(text):
            if not seg:
                continue
            bold_md     = re.fullmatch(r'\*\*([^*]+)\*\*', seg)
            kaiti_md    = re.fullmatch(r'\*([^*]+)\*', seg)
            bold_html   = re.fullmatch(r'<(?:b|strong)>([^<]*)</(?:b|strong)>', seg, re.IGNORECASE)
            em_html     = re.fullmatch(r'<em>([^<]*)</em>', seg, re.IGNORECASE)
            italic_html = re.fullmatch(r'<i>([^<]*)</i>', seg, re.IGNORECASE)
            br_tag      = re.fullmatch(r'<br\s*/?>', seg, re.IGNORECASE)
            span_m      = re.fullmatch(r'<span\b([^>]*)>(.*?)</span>', seg, re.IGNORECASE | re.DOTALL)
            fn_sup_m    = re.search(r'<a\b[^>]*>([^<]*)</a>', seg) if seg.startswith('<sup') else None
            fn_sup      = fn_sup_m if (fn_sup_m and fn_sup_m.group(1).strip().isdigit()) else None
            fn_m        = re.fullmatch(r'\[\^(\d+)\]', seg)
            html_tag    = re.fullmatch(r'<[^>]+>', seg)

            if bold_md:
                run = paragraph.add_run(bold_md.group(1))
                self._apply_font(run, ascii_font, east_font, size=size, bold=True)
            elif kaiti_md:
                run = paragraph.add_run(kaiti_md.group(1))
                self._apply_font(run, ascii_font, '標楷體', size=size)
            elif bold_html:
                run = paragraph.add_run(bold_html.group(1))
                self._apply_font(run, ascii_font, east_font, size=size, bold=True)
            elif em_html:
                run = paragraph.add_run(em_html.group(1))
                self._apply_font(run, ascii_font, '標楷體', size=size)
            elif italic_html:
                run = paragraph.add_run(italic_html.group(1))
                self._apply_font(run, ascii_font, east_font, size=size, italic=True)
            elif br_tag:
                # 段落內換行（w:br）
                br_run = paragraph.add_run()
                br_run._element.append(OxmlElement('w:br'))
            elif span_m:
                attrs   = span_m.group(1)
                inner   = re.sub(r'<[^>]+>', '', span_m.group(2)).strip()
                style_v = ''
                class_v = ''
                sm = re.search(r'style=["\']([^"\']*)["\']', attrs, re.IGNORECASE)
                cm = re.search(r'class=["\']([^"\']*)["\']', attrs, re.IGNORECASE)
                if sm:
                    style_v = sm.group(1).lower().replace(' ', '')
                if cm:
                    class_v = cm.group(1).lower()
                if not inner:
                    pass
                elif 'kaiti' in class_v:
                    # <span class="kaiti"> → 標楷體
                    run = paragraph.add_run(inner)
                    self._apply_font(run, ascii_font, '標楷體', size=size)
                else:
                    # 解析 style 屬性中的 color 與 font-size
                    color_m = re.search(r'color:\s*(#[0-9a-fA-F]{3,6}|\w+)', style_v)
                    size_m  = re.search(r'font-size:\s*(\d+(?:\.\d+)?)pt', style_v)
                    if color_m or size_m or 'font-size:1rem' in style_v or 'font-size:0.' in style_v:
                        sz = int(size_m.group(1)) if size_m else (10 if ('font-size:1rem' in style_v or 'font-size:0.' in style_v) else 12)
                        run = paragraph.add_run(inner)
                        self._apply_font(run, ascii_font, east_font, size=sz)
                        if color_m:
                            try:
                                run.font.color.rgb = RGBColor(*self._hex_rgb(color_m.group(1)))
                            except Exception:
                                pass
                    else:
                        # 其他 span → 以正常字型輸出
                        run = paragraph.add_run(inner)
                        self._apply_font(run, ascii_font, east_font, size=size)
            elif fn_sup:
                self._add_footnote_ref(paragraph, fn_sup.group(1))
            elif fn_m:
                self._add_footnote_ref(paragraph, fn_m.group(1))
            elif html_tag:
                pass  # 略過其他 HTML 標籤（不輸出原始文字）
            else:
                # 解碼 HTML 實體（&nbsp; &#9679; 等）
                seg = _html_mod.unescape(seg)
                # 折疊編輯器在 ** 前後注入的多餘空格（網頁排版用，Word 不需要）
                seg = re.sub(r' {2,}', ' ', seg)
                # 拆分 emoji，用 Segoe UI Symbol；非 emoji 再拆特殊字元
                for chunk, is_emoji in self._split_emoji(seg):
                    if is_emoji:
                        run = paragraph.add_run(chunk)
                        self._apply_font(run, 'Segoe UI Symbol', 'Segoe UI Symbol', size=size)
                    else:
                        # 拆分特殊字元 ‧／∕ → 新細明體（PMingLiU）
                        for sub, is_special in self._split_special_chars(chunk):
                            r = paragraph.add_run(sub)
                            ef = 'PMingLiU' if is_special else east_font
                            self._apply_font(r, ascii_font, ef, size=size)

    # ── 頁首頁尾 ────────────────────────────────────────────

    def add_header_footer(self, issue_number, issue_title, article_id, article_title):
        section = self.doc.sections[0]

        # 啟用奇偶頁不同
        sectPr = section._sectPr
        sectPr.insert(0, OxmlElement('w:evenAndOddHeaders'))

        dark = self._hex_rgb('#262626')

        # 從 article_id 提取文章序號（"7-4" → "04"）
        id_str = str(article_id) if article_id else '04'
        num_m  = re.search(r'-(\d+)', id_str)
        display_num = num_m.group(1).zfill(2) if num_m else id_str.zfill(2)

        # ── 偶數頁（左頁）頁首 ──
        even_hdr = section.even_page_header
        for p in even_hdr.paragraphs: p.clear()

        p1 = even_hdr.paragraphs[0]
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after  = Pt(0)
        r1 = p1.add_run(f'無境界者｜Vol. {issue_number}（2026.02-04）')
        self._apply_font(r1, 'Times New Roman', '微軟正黑體', size=10, color=dark)

        p2 = even_hdr.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(0)
        # 用 tab stop 讓底線延伸到第一行寬度再多 6 格
        # 微軟正黑體 10pt：第一行約 3200 twips，再加 6 半形空格（≈600 twips）= 3800
        pPr2 = p2._element.get_or_add_pPr()
        tabs2 = OxmlElement('w:tabs')
        tab2  = OxmlElement('w:tab')
        tab2.set(qn('w:val'), 'left')
        tab2.set(qn('w:pos'), '3800')   # twips，可依實際效果微調
        tabs2.append(tab2)
        pPr2.append(tabs2)
        r2 = p2.add_run(issue_title + '\t')   # tab 延伸底線到 tab stop
        self._apply_font(r2, 'Times New Roman', '微軟正黑體', size=10, color=dark)
        r2.underline = True

        # ── 奇數頁（右頁）頁首 ──
        # 第一行空白佔位（與偶數頁第一行等高），第二行才是實際文字
        # 讓兩者底部對齊在同一高度
        odd_hdr = section.header
        for p in odd_hdr.paragraphs: p.clear()

        p3_blank = odd_hdr.paragraphs[0]
        p3_blank.paragraph_format.space_before = Pt(0)
        p3_blank.paragraph_format.space_after  = Pt(0)
        r3_blank = p3_blank.add_run('')
        self._apply_font(r3_blank, 'Times New Roman', '微軟正黑體', size=10, color=dark)

        p3 = odd_hdr.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after  = Pt(0)
        # 奇數頁：前6空白 + 編號 + 標題 + 後1空格，靠右對齊，底線包含後面空格
        header_text = f'      {display_num} {article_title} '   # 後1個半形空格
        r3 = p3.add_run(header_text)
        self._apply_font(r3, 'Times New Roman', '微軟正黑體', size=10, color=dark)
        # 底線需包住尾端空格：改用段落底框線模擬，或直接設 underline
        # Word 會裁掉行尾空格的底線；改加一個零寬無破壞空格確保保留
        r3.underline = True
        # 加一個不換行空格（U+00A0）確保尾端有底線
        r3_end = p3.add_run('\u00a0')
        self._apply_font(r3_end, 'Times New Roman', '微軟正黑體', size=10, color=dark)
        r3_end.underline = True

        # ── 偶數頁（左頁）頁尾：空行 + 頁碼靠左（10pt）──
        even_ftr = section.even_page_footer
        for p in even_ftr.paragraphs: p.clear()
        efp_blank = even_ftr.paragraphs[0]
        efp_blank.paragraph_format.space_before = Pt(0)
        efp_blank.paragraph_format.space_after  = Pt(0)
        efp = even_ftr.add_paragraph()
        efp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        efp.paragraph_format.space_before = Pt(0)
        efp.paragraph_format.space_after  = Pt(0)
        self._insert_page_number(efp, size=10)

        # ── 奇數頁（右頁）頁尾：空行 + 頁碼靠右（10pt）──
        odd_ftr = section.footer
        for p in odd_ftr.paragraphs: p.clear()
        ofp_blank = odd_ftr.paragraphs[0]
        ofp_blank.paragraph_format.space_before = Pt(0)
        ofp_blank.paragraph_format.space_after  = Pt(0)
        ofp = odd_ftr.add_paragraph()
        ofp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ofp.paragraph_format.space_before = Pt(0)
        ofp.paragraph_format.space_after  = Pt(0)
        self._insert_page_number(ofp, size=10)

    def _insert_page_number(self, paragraph, size=10):
        run = paragraph.add_run()
        rPr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz');   sz.set(qn('w:val'), str(size * 2))
        szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), str(size * 2))
        rPr.append(sz); rPr.append(szCs)
        run._element.insert(0, rPr)
        fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin')
        ins = OxmlElement('w:instrText')
        ins.set(qn('xml:space'), 'preserve'); ins.text = ' PAGE '
        fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'end')
        run._element.append(fc1)
        run._element.append(ins)
        run._element.append(fc2)

    def save(self, filename):
        self._finalize_footnotes()
        self.doc.save(filename)
        print(f'✅ Word 文件已生成：{filename}')


# ════════════════════════════════════════
# 主程式
# ════════════════════════════════════════

def generate_article_docx(article_data, output_path):
    generator = ProfessionalDocxGenerator()
    generator.set_footnotes(article_data.get('footnotes', []))

    is_toc = (article_data.get('article_type') == 'toc' or
              article_data.get('title') in ('目次', '目錄'))

    category_colors = {
        '封面故事': '#7d6c29', '專題文章': '#C00000',
        '人物專訪': '#FFC000', '評論與回應': '#ED7D31',
        '生命故事': '#46b175', '公告與剪影': '#6a5acd',
        '時事評論': '#0070C0', '時事感想': '#0070C0',
        '文藝創作': '#27408b', '光影時刻': '#7d6c29',
        '實驗園地': '#db7093', '文獻與翻譯': '#548235',
    }
    category = article_data.get('category', '')
    if category:
        generator.add_category_tag(category, category_colors.get(category, '#000000'))
    else:
        # 無欄目時補同高空行（14pt），確保所有文章標題垂直位置一致
        p = generator.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run('\u00a0')   # non-breaking space 撐開行高
        run.font.size = Pt(14)

    if article_data.get('title'):
        generator.add_title(article_data['title'])
    if article_data.get('subtitle'):
        generator.add_subtitle(article_data['subtitle'])

    generator.add_decoration_line()
    blank = generator._add_blank_line()   # 空行（裝飾線與作者之間）
    run = blank.add_run('\u00a0')
    run.font.size = Pt(14)

    generator.add_author(
        article_data.get('author'),
        article_data.get('author_title'),
        article_data.get('remark'),
    )
    if not is_toc:
        generator._add_blank_line()   # 空行（作者與關鍵字之間）
        generator._add_blank_line()

    if article_data.get('keyword'):
        generator.add_keywords(article_data['keyword'])
        generator._add_blank_line()
        generator._add_blank_line()

    if article_data.get('content'):
        generator.add_content(article_data['content'])

    if is_toc:
        generator.save(output_path)
        return

    generator.add_header_footer(
        article_data.get('issue', 7),
        article_data.get('issue_title', '火燒島上的《耶穌傳》'),
        article_data.get('id', '04'),
        article_data.get('title', '無標題'),
    )

    generator.save(output_path)


# ════════════════════════════════════════
# 命令列執行
# ════════════════════════════════════════

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('使用方式：python generate_docx.py <article.json> <output.docx>')
        sys.exit(1)

    input_json  = sys.argv[1]
    output_docx = sys.argv[2]

    try:
        with open(input_json, 'r', encoding='utf-8') as f:
            article_data = json.load(f)
        generate_article_docx(article_data, output_docx)

    except FileNotFoundError:
        print(f'❌ 找不到檔案：{input_json}'); sys.exit(1)
    except json.JSONDecodeError as e:
        print(f'❌ JSON 格式錯誤：{e}'); sys.exit(1)
    except Exception as e:
        print(f'❌ 發生錯誤：{e}')
        import traceback; traceback.print_exc()
        sys.exit(1)
