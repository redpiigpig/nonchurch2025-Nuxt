# -*- coding: utf-8 -*-
"""
Part 2: remaining body fixes + footnote fixes via ZIP XML.
"""
import sys, os, io, shutil, zipfile, tempfile
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from lxml import etree

BASE = r"c:\Users\user\Desktop\nonchurch-nuxt\stores\無境界者雜誌\07-第七期"

FILES = {
    "7-5":  "7-5聆聽被遺忘的苦難.docx",
    "7-6":  "7-6In是Siáng？.docx",
    "7-7":  "7-7被國家剝奪權利的人.docx",
    "7-8":  "7-8黑暗中的那點光.docx",
    "7-9":  "7-9跨地域的緣與情.docx",
    "7-14": "7-14蔡依林有沒有在拜撒旦到底干那些基督徒什麼事？.docx",
}

BODY_CORRECTIONS = [
    # ── 7-5 ──────────────────────────────────────────────────────────────
    ("7-5", "生出面對苦難",     "孕育出面對苦難"),
    ("7-5", "去成都還不知道",   "去成，都還不知道"),
    ("7-5", "與和丈夫同一國",   "和丈夫同一國"),
    ("7-5", "西伯利亞戰友們",   "西伯利亞的戰友們"),

    # ── 7-6 ──────────────────────────────────────────────────────────────
    ("7-6", "讀者讀者",         "讀者"),

    # ── 7-7 ──────────────────────────────────────────────────────────────
    ("7-7", "Just-world fallacy",   "Just-world theory"),
    ("7-7", "被一同被押上車",       "被一同押上車"),
    ("7-7", "積極的活出",           "積極地活出"),
    ("7-7", "直接了當地",           "直截了當地"),

    # ── 7-8 ──────────────────────────────────────────────────────────────
    ("7-8", "他們願意陪著",     "他/她願意陪著"),
    ("7-8", "沈默",             "沉默"),
    ("7-8", "沈重",             "沉重"),

    # ── 7-9 ──────────────────────────────────────────────────────────────
    ("7-9", "是指著矢内原",       "是指矢内原"),
    ("7-9", "留下印象的是，則是", "留下印象的，則是"),
    ("7-9", "不僅陳茂棠",         "不僅是陳茂棠"),
    ("7-9", "忽然之間會固定早退", "忽然開始固定早退"),
    ("7-9", "規劃名為",           "歸化名為"),
    ("7-9", "被譽爲",             "被譽為"),
    ("7-9", "進而，以下將談及",   "接下來，以下將談及"),
    ("7-9", "足立並新渡戸",       "足立與新渡戶"),
    ("7-9", "Yanaibara",          "Yanaihara"),
    ("7-9", "Fujjida",            "Fujita"),
    ("7-9", "統治時期間",         "統治時期"),

    # ── 7-14 ─────────────────────────────────────────────────────────────
    ("7-14", "整理其教義成六大提倡", "並整理其教義，將其化為六大提倡"),
]

# (file_key, footnote_w:id, old_text, new_text)
FOOTNOTE_CORRECTIONS = [
    ("7-7",  8, "以預防他如果在大馬被逮捕的話", "以防他若在大馬被逮捕"),
    ("7-7",  9, "J. Glenn Beall. Jr",           "J. Glenn Beall, Jr."),
    ("7-11", 4, "一切事務",                      "一切事物"),
]


# ─────────────────────────────────────────────────────────────────────────────
# Body replacement helpers
# ─────────────────────────────────────────────────────────────────────────────

def replace_in_para(para, old, new):
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return False
    idx = full.index(old)
    end = idx + len(old)
    pos = 0
    runs = para.runs
    run_spans = []
    for run in runs:
        run_spans.append((pos, pos + len(run.text)))
        pos += len(run.text)
    a_start = next((i for i, (s, e) in enumerate(run_spans) if s <= idx < e), None)
    a_end   = next((i for i, (s, e) in enumerate(run_spans) if s < end <= e), None)
    if a_start is None or a_end is None:
        return False
    pre  = full[run_spans[a_start][0]:idx]
    post = full[end:run_spans[a_end][1]]
    runs[a_start].text = pre + new + post
    for i in range(a_start + 1, a_end + 1):
        runs[i].text = ""
    return True


def replace_in_doc_body(doc, old, new):
    count = 0
    for para in doc.paragraphs:
        if replace_in_para(para, old, new):
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replace_in_para(para, old, new):
                        count += 1
    return count


# ─────────────────────────────────────────────────────────────────────────────
# Footnote fix via ZIP XML manipulation
# ─────────────────────────────────────────────────────────────────────────────

def fix_footnote_in_docx(fpath, fn_id_int, old, new):
    """
    Directly modifies word/footnotes.xml inside the docx zip.
    Returns number of replacements made.
    """
    # Read original zip
    with open(fpath, 'rb') as f:
        original = f.read()

    changed = 0
    buf = io.BytesIO()

    with zipfile.ZipFile(io.BytesIO(original), 'r') as zin, \
         zipfile.ZipFile(buf, 'w', compression=zipfile.ZIP_DEFLATED) as zout:

        for item in zin.infolist():
            item_name = item.filename if hasattr(item, 'filename') else item.name
            data = zin.read(item_name)

            if item_name == 'word/footnotes.xml':
                root = etree.fromstring(data)
                ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                W = f'{{{ns}}}'

                for fn_el in root.iter(f'{W}footnote'):
                    wid = fn_el.get(f'{W}id')
                    if wid is None or int(wid) != fn_id_int:
                        continue
                    for t_el in fn_el.iter(f'{W}t'):
                        if t_el.text and old in t_el.text:
                            t_el.text = t_el.text.replace(old, new)
                            changed += 1

                data = etree.tostring(root, xml_declaration=True,
                                      encoding='UTF-8', standalone=True)

            zout.writestr(item_name, data)

    if changed:
        with open(fpath, 'wb') as f:
            f.write(buf.getvalue())

    return changed


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def main():
    by_file = {}
    for (key, old, new) in BODY_CORRECTIONS:
        by_file.setdefault(key, []).append((old, new))

    fn_by_file = {}
    for (key, fn_id, old, new) in FOOTNOTE_CORRECTIONS:
        fn_by_file.setdefault(key, []).append((fn_id, old, new))

    all_keys = sorted(set(list(by_file.keys()) + list(fn_by_file.keys())),
                      key=lambda k: (int(k.split('-')[1]), k))

    for key in all_keys:
        fname = FILES.get(key)
        if not fname:
            print(f"[SKIP] {key}: no filename")
            continue
        fpath = os.path.join(BASE, fname)
        if not os.path.exists(fpath):
            print(f"[SKIP] {key}: file not found")
            continue

        print(f"\n{'='*60}")
        print(f"  {key}: {fname}")
        print(f"{'='*60}")

        # Body corrections
        body_changes = by_file.get(key, [])
        if body_changes:
            try:
                doc = Document(fpath)
                changed = False
                for (old, new) in body_changes:
                    n = replace_in_doc_body(doc, old, new)
                    if n > 0:
                        print(f"  OK [{n}x] {old!r} -> {new!r}")
                        changed = True
                    else:
                        print(f"  !! NOT FOUND: {old!r}")
                if changed:
                    doc.save(fpath)
                    print(f"  >> Body saved.")
                else:
                    print(f"  (body: no changes)")
            except PermissionError:
                print(f"  [LOCKED] {fname} is open — close it and re-run.")
                continue

        # Footnote corrections (via ZIP XML)
        for (fn_id, old, new) in fn_by_file.get(key, []):
            try:
                n = fix_footnote_in_docx(fpath, fn_id, old, new)
                if n > 0:
                    print(f"  OK fn{fn_id} [{n}x] {old!r} -> {new!r}")
                else:
                    print(f"  !! fn{fn_id} NOT FOUND: {old!r}")
            except PermissionError:
                print(f"  [LOCKED] cannot fix footnote — file is open.")

    print("\nDone.")


if __name__ == "__main__":
    main()
